import random

leap1=[('I totally agree with this opinion..','私たちはこの意見に全面的に賛成です。〔No.1〕'),

       ('strongly oppose the plan.','その計画に強固に反対する。〔No.2〕'),

       ('advise him not to eat too much.','食べ過ぎないよう彼に忠告する。〔No.3〕'),

       ('tips on how to use the library.','図書館の利用法に関するヒント。〔No.4〕'),

       ('discuss the problem with a specialist.','専門家とその問題について話し合う。〔No.5〕'),

       ("blame the car's brakes for the accident.",'その事故の原因は車のブレーキにあるとする。〔No.6〕'),

       ('argue that reading aloud is important.','音読は大切だと主張する 。〔No.7〕'),

       (' claim that a vegetarian diet is better than meat diet.','菜食は肉食より優れていると主張する。〔No.8〕'),

       (' complain about their loud music.','彼らのうるさい音楽に文句を言う 。〔No.9〕'),

       ('offer him some coffee.','彼にコーヒーはどうですかと尋ねる(申し出る)。〔No.10〕'),

       ('hat letter suggests that she is quite happy in her job.','その手紙は彼女が仕事をかなり楽しんでいることを示唆している。〔No.11〕'),

       ('What would you recommend?','＜レストランなどで＞何がお勧めですか(何を勧めますか)。〔No.12〕'),

       (' I am grateful for your help.','ご助力に感謝しています。〔No.13〕'),

       ('apologize to her for being late.','遅刻したことを彼女に謝る。〔No.14〕'),

       ("Don't make excuses.",'言い訳をするな。〔No.15〕'),

       ('celebrate her 18th birthday.','彼女の18歳の誕生日を祝う。〔No.16〕'),

       ('admire him for his great performance.','すばらしい演技に対して彼を称賛する。〔No.17〕'),

       ('I was deeply impressed by his speech.','彼の演説に深い感銘を受けた(感銘を与えられた) 。〔No.18〕'),

       ('an award ceremony.','授賞式。〔No.19〕'),

       ('Could you describe your lost bag?','なくしたバッグ(の特徴)を説明していただけませんか。〔No.20〕'),

       (' explain the rules of baseball to him.','彼に野球のルールを説明する。〔No.21〕'),

       ('communicate with each other in sign language.','手話でお互いに意思の疎通をはかる。〔No.22〕'),

       ('express your opinions clearly.','はっきりと意見を言い(を表現し)なさい。〔No.23〕'),

       (' make a promise to lose weight.','減量するという約束をする。〔No.24〕'),

       ('various kinds of information.','さまざまな情報。〔No.25〕'),

       ('science and technology.','科学技術(科学と技術)。〔No.26〕'),

       ('do research on space rockets.','宇宙ロケットに関する研究を行う。〔No.27〕'),

       ('collect materials for a paper.','レポートの材料を集める。〔No.28〕'),

       (' artificial intelligence.','人工知能。〔No.29〕'),

       ('an electric car.','電気自動車。〔No.30〕'),

       (' invent a speaking robot.','会話のできるロボットを発明する。〔No.31〕'),

       ('discover a human mummy.','人間のミイラを発見する。〔No.32〕'),

       ('develop into a big city.','大都市に発展する。〔No.33〕'),
       
       ('improve my tennis skills.','テニスの技術を高める。〔No.34〕'),

       ('improve my ability to speak English.','英語を話す力を伸ばす。〔No.35〕'),

       ('how my musical talent.','音楽の才能を発揮する。〔No.36〕'),

       ('make an effort to be on time.','時間に間に合うように努力する。〔No.37〕'),

       ('We have practice on Saturday.','土曜日に練習がある。〔No.38〕'),

       ('a soccer game between Japan and Brazil.','日本とブラジルのサッカーの試合。〔No.39〕'),
       
       ('achieve his goal of becoming a vet.','獣医になるという目標を達成する。〔No.40〕'),

       ('I managed to catch the last train.','何とか終電に乗ることができた。〔No.41〕'),

       (' improve my cooking skills.','私の料理の腕を上達させる 。〔No.42〕'),

       ('produce rice and vegetables.','米と野菜を生産する。〔No.43〕'),

       ('create a website.','ホームページを創る。〔No.44〕'),

       (' establish a close friendship with him.','彼と親密な友情を築く。〔No.45〕'),

       ('This will save you a lot of trouble.','これによって多くの手間が省ける。〔No.46〕'),

       ('take the medicine twice a day.','１日に２回その薬を飲む。〔No.47〕'),

       ('The patient is getting better.','その患者は快方に向かっている。〔No.48〕'),

       ('My body is in good condition.','私の体は調子がよい。〔No.49〕'),
       
       ('medical expenses.','医療費 。〔No.50〕'),

       ('have a lot of stress.','ストレスが多い。〔No.51〕'),

       ('I am sufferam ing from jet lag.','時差ぼけに苦しんでいる。〔No.52〕'),

       ('get light exercise.','軽い運動をする。〔No.53〕'),

       (' breathe deeply.','深呼吸する。〔No.54〕'),

       ("I'm really thirsty.",'本当にのどが渇いた。〔No.55〕'),

       ('physical health.','体の健康。〔No.56〕'),

       ('I have a fever.','熱がある。〔No.57〕'),

       ('build up my physical strength.','体力をつける。〔No.58〕'),
       
       (' was moved to tears.','私は感動して泣いた。〔No.59〕'),

       ('This apple tastes sweet.','このリンゴは甘い(味がする)。〔No.60〕'),

       ('It is against the rules.','それは規則違反だ。〔No.61〕'),

       (' play an important role in society.','社会で重要な役割を果たす。〔No.62〕'),

       (' have a habit of making excuses.','言い訳をする癖がある。〔No.63〕'),

       ('Different countries have different customs.','習慣は国よっても違う(異なる国は異なる習慣をもつ)。〔No.64〕'),

       ('Japanese traditions.','日本の伝統。〔No.65〕'),

       ("women's status in society.",'社会での女性の地位。〔No.66〕'),

       ('Parking here is against the law.','ここでの駐車は法律違反だ。〔No.67〕'),

       ('My ancestors were French.','私の祖先はフランス人だった。〔No.68〕'),

       ('The population of Tokyo is larger than that of Osaka.','東京の人口は大阪の人口より多い。〔No.69〕'),

       ('his native language.','彼の母(国)語。〔No.70〕'),

       ('find a job abroad.','海外で仕事を見つける。〔No.71〕'),

       ('a local bank.','地元の銀行。〔No.72〕'),

       (' according to a recent survey.','最近の調査によると。〔No.73〕'),

       ('a body of strong evidence.','一連の有力な証拠。〔No.74〕'),

       ('he value of the painting.','その絵の価値。〔No.75〕'),

       ('a national treasure.','国宝。〔No.76〕'),

       (' follow the latest fashion.','最新の流行を追う。〔No.77〕'),

       ('the general public.','一般大衆。〔No.78〕'),

       ('receive the 45% of the vote.','投票数の45%を獲得する。〔No.79〕'),

       ('the Japanese government.','日本政府。〔No.80〕'),

       (' the most powerful nation in the world.','世界で一番力のある国。〔No.81〕'),

       ('What is the capital of Canada?.','カナダの首都はどこですか。〔No.82〕'),

       (' the United states of America.','アメリカ合衆国 。〔No.83〕'),

       ('the political situation in Russia.','ロシアの政治情勢。〔No.84〕'),

       ('pay a very high price for the painting.','その絵に高額(の価格)を支払う。〔No.85〕'),

       ('The suit cost me 500 dollars.','そのスーツは500ドルした(かかった)。〔No.86〕'),

       ('pay a large sum for antiques.','骨董品に大金を使う。〔No.87〕'),

       ('our budget for this trip.','今回の旅の予算。〔No.88〕'),

       ('pay a cancellation fee.','キャンセル料を払う。〔No.89〕'),

       ('Taxi fares are low in Singapore..','シンガポールのタクシー料金(運賃)は安い。〔No.90〕'),

       ('May we have the bill?','bill。〔No.91〕'),

       ("Japan's trade with the U.S.",'日本の対米貿易。〔No.92〕'),

       ('achieve wealth and power.','富と権力を手にす。〔No.93〕'),

       (' boost the Japanese economy.','日本経済を強化する。〔No.94〕'),

       ('work for a big company.','大きな会社で働く。〔No.95〕'),

       ('accomplish a difficult task.','困難な仕事を成し遂げる。〔No.96〕'),

       ('earn money to pay my school fees.','学費を払う金を稼ぐ。〔No.97〕'),

       ('My back hurts a lot.','腰がとても痛い。〔No.98〕'),

       ('injure my back.','腰を痛める。〔No.99〕'),

       ('seriously damage the environment.','環境に深刻な被害を与える。〔No.100〕'),

       (' destroy all weapons.','すべての武器を破壊する。〔No.101〕'),

       ('You ruined my plan.','君のせいで計画が台無しだよ。〔No.102〕'),

       ('Tigers are in danger of extinction.','トラが絶滅の危機にある。〔No.103〕'),

       (' chat with a neighbor.','近所の人と話す。〔No.104〕'),

       ('draw a large audience.','大観衆を集める。〔No.105〕'),

       ('A crowd rushed into the hall.','群衆が会場になだれ込んだ。〔No.106〕'),

       ('the author of War and Peace.','『戦争と平和』の著者。〔No.107〕'),

       (' Most of the staff is from Australia.','職員の大半はオーストラリア出身です。〔No.108〕'),

       (' the clerks at that department store.','あのデパートの店員。〔No.109〕'),

       ('The restaurant was crowded with customers from abroad.','そのレストランは外国人客で混雑していた。〔No.110〕'),

        ('There were a lot of passengers on the train.','列車には多くの乗客が載っていた。〔No.111〕'),

        ('give up my seat to elderly people.','お年寄り(年配の人)に席を譲る。〔No.112〕'),

        (' female workers.','女性の労働者。〔No.113〕'),

        ('head for the beach.','海辺へ向かう。〔No.114〕'),

        ('Someone was following me.','誰かが私をつけていた。〔No.115〕'),

        ('wander around the town.','街をぶらぶら歩き回る。〔No.116〕'),

        ( 'travel around the world.','世界中を旅する。〔No.117〕'),

        ('pass the post office.','郵便局(の前)を通り過ぎる。〔No.118〕'),

        ('Things are likely to improve.','事態はよくなりそうだ。〔No.119〕'),

        ('t is only natural that you cannot catch words you do not know.','知らない単語を聞き取れないのはごく当然のことだ。〔No.120〕'),

        ('He is certain to become a star player.','彼はきっとスター選手になるよ。〔No.121〕'),

        ('The repairs will probably take a month.','修理はおそらく1か月はかかるだろう。〔No.122〕'),

        ('We are faced with a crisis.','私たちは危機に直面している。〔No.123〕'),

        ('avoid making eye contact with monkeys.','サルとは目を合わせないようにしなさい。〔No.124〕'),

        ('solve the math problem.','その数学の問題を解く 。〔No.125〕'),

        (' have trouble sleeping.','寝付けない(眠るのに苦労する)。〔No.126〕'),

        ('an international issue.','国際問題。〔No.127〕'),

        ('cheer all the players on.','すべての選手たちに声援を送る。〔No.128〕'),

        (' encourage people to bike to work.','人々に自転車通勤するよう促す。〔No.129〕'),

        (' firmly support her.','強く彼女を支持する。〔No.130〕'),

        ('The heavy snow prevented us from going out.','大雪のため外出できなかった。〔No.131〕'),

        (' deny that he is a great scientist.','彼が偉大な科学者であることを否定する 。〔No.132〕'),

        ('he Internet has enabled us to receive information right away.','インターネットのおかげで情報をすぐに得られるようになった。〔No.133〕'),

        ('succeed in landing on the moon.','月面着陸に成功する。〔No.134〕'),

        ('miss the last train.','終電を逃す。〔No.135〕'),

        (' fail in business.','商売で失敗する 。〔No.136〕'),

        ('make grammatical mistakes.','文法ミスをする。〔No.137〕'),

        ('check the brakes.','そのブレーキを調べる。〔No.138〕'),

        ('hide the comic under the pillow.','枕の下に漫画を隠す。〔No.139〕'),

        ('draw her portrait.','彼女の似顔絵を描く。〔No.140〕'),

        ('join the baseball team.','野球部に入る。〔No.141〕'),

        ('throw me a rope.','ロープを投げてくれ。〔No.142〕'),

        (' operate five machines at once.','１度に5台の機械を操作する。〔No.143〕'),

        ('repair a bicycle.','自転車を修理する。〔No.144〕'),

        ('She ews all her own dresses.','彼女は自分の服を自分で縫っている。〔No.145〕'),

        ('raise your hand.','手を上げて。〔No.146〕'),

        ('serve wine to the guests.','客にワインを出す。〔No.147〕'),

        (' pour tea for all of the guests.','客みんなにお茶を注ぐ 。〔No.148〕'),

        ('spill coffee all over my laptop.','コーヒーをノートパソコンの上にこぼす。〔No.149〕'),

        ('pretend to be sleeping.','寝ているふりをする。〔No.150〕'),

        ('tell him to behave himself.','彼に行儀よくふるまうように言う。〔No.151〕'),

        ('cannot bear the noise.','騒音に耐えられない。〔No.152〕'),

        ('explore the surface of Mars.','火星の地表を探査する。〔No.153〕'),

        ('What happened to her?','彼女に何が起こったの。〔No.154〕'),

        ('The singer appeared on the stage.','その歌手がステージに現れた。〔No.155〕'),

        ('The cause of the accident remains unknown.','その事故の原因はいまだ知られていない。〔No.156〕'),

        (' eat grass in order to survive.','生き延びるため草を食べる。〔No.157〕'),

        ('Sam belongs to the Sales Department.','サムは営業部に所属している。〔No.158〕'),

        ('represent Japan at the World Cup.','ワールドカップ日本代表となる。〔No.159〕'),

        ('This novel is based on facts.','この小説は事実に基づいている。〔No.160〕'),

        ('The price includes postage.','価格には送料が含まれています。〔No.161〕'),

        ('his soup contains various herbs.','このスープにはさまざまな薬草が入っている。〔No.162〕'),

        ('with my own eyes.','私自身の目で。〔No.163〕'),

        (' share a table with a stranger.','知らない人と相席(席を共有)する。〔No.164〕'),

        ('collect teddy bears.','クマのぬいぐるみを集める。〔No.165〕'),

        ('gather information.','情報を集める。〔No.166〕'),

        ('This temple has a history of a thousand years.','この寺には1000年の歴史がある。〔No.167〕'),

        (' a required subject.','必須科目 。〔No.168〕'),

        ('eceive a good education.','よい教育を受ける。〔No.169〕'),

        ('have a good knowledge of cinema.','映画に関する知識が豊富だ。〔No.170〕'),

        ('wear a school uniform.','学校の制服を着ている。〔No.171〕'),

        ('She is in the eighth grade.','彼女は８年生(中学２年生)だ。〔No.172〕'),

        ('a college senior.','大学４年生 。〔No.173〕'),

        (' graduate from high school.','高校を卒業する。〔No.174〕'),

        ('decide to live on my own.','一人で暮らすことを決心する。〔No.175〕'),

        (' judge someone by their appearance.','人をが意見で判断する。〔No.176〕'),

        ('quit my job.','仕事をやめる。〔No.177〕'),

        (' retire from the company.','会社を(定年で)退職する。〔No.178〕'),

        (' a review of the tax system.','税制の見直し。〔No.179〕'),

        ('choose what to wear to school.','学校へ着ていく服を選ぶ 。〔No.180〕'),

        ('imagine life without electricity.','電気のない生活を想像してごらん。〔No.181〕'),

        ('guess what is in this bag.','この袋に何が入っているかを当ててごらん。〔No.182〕'),

        ('We are expecting a typhoon.','台風が来るらしい。〔No.183〕'),

        (' predict when cherry blossoms will start to flower.','桜の開花時期を予測する。〔No.184〕'),

        ('I remember meeting her here.','彼女にここで会ったことを覚えています。〔No.185〕'),

        ('This song reminds me of him.','この歌は彼のことを私に思い出させる。〔No.186〕'),

        ('Jim realized that he did not love her any more.','ジムはもう彼女のことを愛していないことに気付いた。〔No.187〕'),

        ('appreciate the importance of discussion.','話し合いの重要性を理解する。〔No.188〕'),

        ("accept other people's opinions.",'他人の意見を受け入れる。〔No.189〕'),

        ('considering that he is single.','彼が独身だということを考慮すると。〔No.190〕'),

        ('Do you mind if I open the window?','窓を開けてもいいですか(窓を開けるのを気にしますか)。〔No.191〕'),

        (' I wonder where he has gone.','彼はどこに行ってしまったのだろうか。〔No.192〕'),

        ('It seems that he knows what happened.','彼は起きたことを知っているようだ。〔No.193〕'),

        ('compare city life with country life.','都会の生活を田舎の生活と比較する。〔No.194〕'),

        ('concentrate on my studies.','勉強に集中する。〔No.195〕'),

        ('his book focuses on French culture.','この本はフランス文化に焦点を当てている。〔No.196〕'),

        (' Please allowme to introduce myself.','自己紹介をさせてください。〔No.197〕'),

        ('Tom admitted that he had been mistaken.','トムは自分が間違っていたことを認めた。〔No.198〕'),

        ('deep in the ground.','地中深くに。〔No.199〕'),
 
        ('prevent marine pollution.','海洋汚染を食い止める。〔No.200〕'),

       ('Snow is a rare sight in this area.','雪はこの地域では珍しい光景です。〔No.201〕'),     

("enjoy a night view of Kobe.","神戸の夜景（夜の景色）を楽しむ。〔No.202〕"),

(" the beautiful landscape of Scotland.","スコットランドの美しい風景。〔No.203〕"),

(" prevent nature.","自然を守る。〔No.204〕"),

("sunflower seeds.","ひまわりの種 。〔No.205〕"),

("water a plant.","植物に水をやる。〔No.206〕"),

("Flowers bloomed all over the field.","野原には一面花が咲いていた。〔No.207〕"),

("the rice harvest this year.","今年の米の収穫。〔No.208〕"),

("the chirping of insects.","虫の音。〔No.209〕"),

("There was an earthquake last week.","先週地震が起きた。〔No.210〕"),

("The temperature reached 40 degrees.","気温が40度まで上がった。〔No.211〕"),

("freeze at zero degrees.","ゼロ度で凍る 。〔No.212〕"),

("The flood washed the bridge away.","洪水で橋が流された。〔No.213〕"),

("The pond froze over.","池が一面凍った。〔No.214〕"),

("reflect the sunlight.","日光を反射する。〔No.215〕"),

("a polite young man.","礼儀正しい若い男性。〔No.216〕"),

("It is rude to ask women their age.","女性に年齢を尋ねるのは無礼です。〔No.217〕"),

("live a lonely life.","寂しい生活を送る。〔No.218〕"),

("He is lazy at work.","彼は仕事をいい加減にやる。〔No.219〕"),

("a strict rule.","厳しい規則。〔No.220〕"),

("The Ugly Duckling.","『醜いアヒルの子』。〔No.221〕"),

("intelligent beings.","知的生命体。〔No.222〕"),

("Don't be silly!","ばかなことを言うな！〔No.223〕"),

("get nervous during the exam.","試験中にあがる。〔No.224〕"),

("The noise kept me awake all night.","騒音で一晩中眠れなかった。〔No.225〕"),

("have a positive attitude toward life.","人生に対して前向きな姿勢でいる。〔No.226〕"),

("Printed books have character.","活字の本には味（特徴）がある。〔No.227〕"),

("one of the characteristics of large cities.","大都会の特徴の１つ 。〔No.228〕"),

("Low rainfall is a feature of deserts.","少ない降水量は砂漠の特徴の１つだ。〔No.229〕"),

("check the details of the report.","その報告書の詳細を確認する。〔No.230〕"),

("There are advantages and disadvantages to self-driving cars.","自動運転の車にはメリットとデメリットがある。〔No.231〕"),

("complain about my husband's faults.","夫の欠点について文句を言う。〔No.232〕"),

("improve the quality of life.","生活の質を向上させる。〔No.233〕"),

("Let's see if our answers are correct.","答えが正しいかどうかみてみよう。〔No.234〕"),

("an ideal place for camping.","キャンプにはうってつけの場所。〔No.235〕"),

("It's not fair!","それは不公平だ。〔No.236〕"),

("clothes appropriate for a job interview.","仕事の面接に適切な服。〔No.237〕"),

("Kyoto is famous for its temples.","京都は寺で有名だ。〔No.238〕"),

("when I was in elementary school.","小学生（小学校に在学中）のとき。〔No.239〕"),

("a major cause of global warming.","地球温暖化の主な原因。〔No.240〕"),

("It does not matter whether he succeeds or not.","彼が成功するかどうかは重要ではない。〔No.241〕"),

("The moon is bright tonight.","今夜は月が明るい。〔No.242〕"),

("swim in the brilliant sunshine.","まばゆい太陽光の下で泳ぐ。〔No.243〕"),

("have a lively conversation.","生き生きとした会話をする。〔No.244〕"),

("a comfortable bed.","快適なベッド。〔No.245〕"),

("I find it pleasant to talk with her.","彼女と話すのは楽しいと感じる。〔No.246〕"),

("Call me when it is convenient for you.","都合の良い時に電話してください。〔No.247〕"),

(" false information.","誤った情報。〔No.248〕"),

("I have a terrible sense of direction.","私は方向音痴（ひどい方向感覚）だ。〔No.249〕"),

("an awful smell of cat urine.","ネコの尿のひどい臭い。〔No.250〕"),

("a thin jacket.","薄い上着。〔No.251〕"),

(" tight jeans.","きついジーンズ 。〔No.252〕"),

("wear loose clothes.","ダボっとした（ゆるい）。〔No.253〕"),

("raw fish.","生魚。〔No.254〕"),

("an empty bottle.","空瓶。〔No.255〕"),

("This soap makes your skin very smooth.","この石けんでお肌はツルツル（とても滑らか）になります。〔No.256〕"),

("a direct flight to Paris.","パリへの直行便 。〔No.257〕"),

("Shakespeare is familiar to people all over the world.","シェイクスピアは世界中の人々に知られている。〔No.258〕"),

("She is similar to her mother in appearance.","彼女は見た目が母親に似ている。〔No.259〕"),

("His tastes differ from mine.","彼の趣味は私の（趣味）と異なる。〔No.260〕"),

("The way people greet each other varies from country to country.","挨拶の仕方は国によりさまざまだ。〔No.261〕"),

(" a specific age group.","ある特定の年齢のグループ。〔No.262〕"),

("E-books are becoming common.","電子書籍は普及しつつある。〔No.263〕"),

("These days, it is not unusual for men to wear makeup.","最近は，男性が化粧するのは珍しくない。〔No.264〕"),

("add some milk to the coffee.","コーヒーにミルクを加える。〔No.265〕"),

("increase by 30 percent.","30％増える。〔No.266〕"),

("reduce the amount of garbage.","生ゴミ（の量）を減らす。〔No.267〕"),

("divide the stew into small portions.","シチューを小分けにする。〔No.268〕"),

("count the money.","お金を数える。〔No.269〕"),

("weigh 40 kilograms.","体重が40キロだ。〔No.270〕"),

("a quarter to ten.","10時15分前（9:45)。〔No.271〕"),

("a lot of dolphins.","多くのイルカ。〔No.272〕"),

("a pile of ironed shirts.","アイロンをかけたシャツの山。〔No.273〕"),

("the declining birth rate.","少子化（下がりつつある出生率）。〔No.274〕"),

("memorize a list of figures.","数字の一覧を暗記する。〔No.275〕"),

("I am tired due to (a) lack of sleep.","私は睡眠不足のため疲れている。〔No.276〕"),

("an extra charge.","追加料金。〔No.277〕"),

("I'm not ready yet.","まだ準備ができていません。〔No.278〕"),

("prepare a meal.","食事の準備をする。〔No.279〕"),

("adjust to a new school life.","新しい学校生活に慣れる。〔No.280〕"),

("That rule applies to students.","その規則は学生に適用される。〔No.281〕"),

("What time suits you best?","何時が（あなたに）ご都合がいいですか。〔No.282〕"),

("In some cases, honesty does not pay.","正直が割に合わない場合もある。〔No.283〕"),

("a love scene.","ラブシーン。〔No.284〕"),

(" a chance to play the lead.","主役を演じる機会。〔No.285〕"),

("have an opportunity to speak English.","英語を話す機会を持つ。〔No.286〕"),

("learn from experience.","経験から学ぶ。〔No.287〕"),

("If you have time, help me.","時間があれば手伝ってよ。〔No.288〕"),

("spring has come.","春が来た。〔No.289〕"),

("The station is ten minutes' walk from here.","駅はここから歩いて10分です。〔No.290〕"),

("the Roman period.","ローマ時代。〔No.291〕"),

("a generation gap.","世代間格差。〔No.292〕"),

("celebrate our school's 100th anniversary.","我が校の創立100周年を祝う。〔No.293〕"),

("I moved to this city just recently.","最近この市に引っ越してきました。〔No.294〕"),

("I used to drive to work, but nowadays I take the train.","以前は車で仕事に行っていたが，今は電車を利用する。〔No.295〕"),

("early in the morning.","朝早くに。〔No.296〕"),

("win first prize.","1等賞をとる。〔No.297〕"),

("follow the latest fashions.","最新のファッションを追う。〔No.298〕"),

("in modern society.","現代社会では。〔No.299〕"),

("the latter half of the 20th century.","20世紀後半。〔No.300〕"),

("spend most of my time reading.","時間の大半を読書に費やす。〔No.301〕"),

("Our train was delayed (for) two hours.","私たちが乗った列車が2時間遅れた。〔No.302〕"),

("borrow two books from the library.","図書館から2冊の本を借りる。〔No.303〕"),

("Could you lend me 10,000 yen?.","1万円貸してもらえませんか。〔No.304〕"),

("rent a car for six hours.","6時間，車を借りる。〔No.305〕"),

("provide students with an opportunity to study in the U.K.","イギリス留学の機会を学生に与える。〔No.306〕"),

("on the top shelf.","一番上の棚に。〔No.307〕"),

("drill a hole through a board.","板にドリルで穴を開ける。〔No.308〕"),

("plastic garbage bags.","ビニールのごみ袋 。〔No.309〕"),

("buy sales items in bulk.","特売品をまとめ買いする。〔No.310〕"),

("the front wheel of my bicycle.","自転車の前輪。〔No.311〕"),

(" take notes in class.","授業中にメモをとる。〔No.312〕"),

("open the present.","プレゼントを開ける。〔No.313〕"),

("lift heavy stuff.","重い物を持ち上げる 。〔No.314〕"),

("tie up the box with string.","ひもで箱を縛る。〔No.315〕"),

("enjoy my leisure (time).","余暇を楽しむ。〔No.316〕"),

("a well-balanced diet.","栄養バランスのとれた食事。〔No.317〕"),

("a set of furniture.","家具一式。〔No.318〕"),

("put the leftovers in the refrigerator.","冷蔵庫に残り物を入れる。〔No.319〕"),

("There is little traffic on the street today.","今日は通りに車が少ない（交通量がほとんどない）。〔No.320〕"),

("get caught in a traffic jam.","交通渋滞につかまる 。〔No.321〕"),

("We sold our car for 500,000 yen.","私たちは車を50万円で売った。〔No.322〕"),

("pay five dollars for the book.","その本の代金として5ドルを払う。〔No.323〕"),

("I wear contact lenses.","コンタクトレンズをつけている。〔No.324〕"),

("change my clothes.","服を着替える。〔No.325〕"),

("Will you marry me?","（私と）結婚してくれますか？〔No.326〕"),

("greet each other.","お互いに挨拶をする。〔No.327〕"),

("order rice directly from the farm.","産地に直接米を注文する。〔No.328〕"),

("book a single room for two nights.","シングルルームを2泊予約する。〔No.329〕"),

("a two-hundred-story skyscraper.","200階建ての超高層ビル。〔No.330〕"),

("a dam site.","ダムの用地。〔No.331〕"),

("mow the lawn in our yard.","私たちの庭の芝を刈る。〔No.332〕"),

("the bottom of the sea.","海底。〔No.333〕"),

("stand in line.","列に並ぶ。〔No.334〕"),

("the people sitting in the back row.","後列に座っている人。〔No.335〕"),

("Classical music was playing in the background.","BGM（背景）にクラシック音楽が流れていた。〔No.336〕"),

("the opposite direction.","反対の方向。〔No.337〕"),

("works of modern art.","現代芸術の作品。〔No.338〕"),

("various foreign cultures.","さまざまな外国文化 。〔No.339〕"),

("Japanese animated cartoons.","日本のアニメーション作品（アニメ化された漫画）。〔No.340〕"),

("the plot of this novel.","この小説のあらすじ 。〔No.341〕"),

("play a musical instrument.","楽器を演奏する。〔No.342〕"),

("That is a nice tune.","いい曲だね。〔No.343〕"),

("delete the sentence.","その１文を削除する。〔No.344〕"),

("a magazine article on Portugal.","ポルトガルに関する雑誌の記事。〔No.345〕"),

("a famous passage from the Bible.","聖書の有名な一節。〔No.346〕"),

("increase my vocabulary.","語彙を増やす。〔No.347〕"),

("rumors about him spread throughout the school.","彼のうわさは学校中に広まった。〔No.348〕"),

("How do you spell your name?","名前はどのようにつづりますか。〔No.349〕"),

("How do you pronounce your name?","あなたの名前はどう発音するのですか。〔No.350〕"),

("show signs of recovery.","回復の兆しが見られる。〔No.351〕"),

("What does that mean ?","それはそういう意味ですか。〔No.352〕"),

("This magazine is published twice a year.","この雑誌は年に２回発行されている。〔No.353〕"),

("display the latest fashions.","最新のファッションを展示する。〔No.354〕"),

("You can trust him; he keeps secrets.","彼のことを信頼してよい。彼は秘密を守る。〔No.355〕"),

("Your success depends on how much effort you make.","成功は君の努力（君がどのくらい努力するか）次第だ。〔No.356〕"),

("You can rely on him.","彼なら頼っても大丈夫。〔No.357〕"),

("pray for her safety.","彼女の無事を祈る。〔No.358〕"),

("beg her to overlook my mistakes.","ミスを見逃してもらうよう，彼女にお願いする。〔No.359〕"),

("prefer beef to chicken.","鶏肉より牛肉を好む。〔No.360〕"),

("All America wept.","全米が泣いた。〔No.361〕"),

("hate doing housework.","家事をするのが嫌いだ。〔No.362〕"),

("Don't worry about it.","そんなことを気にするな。〔No.363〕"),

("I'm anxious about her health.","彼女の健康が心配だ。〔No.364〕"),

("I'm not satisfied with the exam result.","試験結果に満足していない。〔No.365〕"),

("It really annoys me when I see people smoking while (they are) walking.","歩きタバコを見ると頭にくる（歩きタバコは私をいらだたせる）。〔No.366〕"),

("I'm sorry to bother you, but could you give me a hand?","ご迷惑をおかけしてすみませんが，手を貸していただけないでしょうか。〔No.367〕"),

("Don't disturb other people.","ほかの人に迷惑をかけないようにしなさい。〔No.368〕"),

("I am frightened of spiders.","クモが怖い（クモにおびえさせられる）。〔No.369〕"),

("The news upset her.","その知らせは彼女を動揺させた。〔No.370〕"),

("I regret that I did not study hard in my teens.","10代のころ一生懸命に勉強しなかったことを後悔している。〔No.371〕"),

("Could you do me a couple of favors?","いくつかお願いがあるのですが。〔No.372〕"),

("have an interest in history.","歴史に興味を持っている。〔No.373〕"),

("It is a pity that you missed the movie.","あなたがその映画を見損ねたのは残念です。〔No.374〕"),

("Our flight was delayed due to the storm.","嵐のために飛行機が遅れた。〔No.375〕"),

("agree with his opinion for the following reasons.","次の理由で彼の意見に賛成だ。〔No.376〕"),

("as a result of inflation.","インフレの結果として。〔No.377〕"),

("Drinking too much has a bad effect on the brain.","過度の飲酒は脳に悪影響を与える。〔No.378〕"),

("have a strong influence on them.","彼らに強い影響を与える。〔No.379〕"),

("cause trouble.","面倒を引き起こす。〔No.380〕"),

("Lack of sleep seriously affects your performance at work.","睡眠不足は仕事（の出来栄え）に悪影響を及ぼす。〔No.381〕"),

("my way of studying.","私の勉強の仕方。〔No.382〕"),

("behave in a responsible manner.","責任ある方法でふるまう。〔No.383〕"),

("What is the purpose of your visit?","あなたの訪問の目的は何ですか。〔No.384〕"),

("for the sake of your health.","健康のため。〔No.385〕"),

("right in front of me.","私のちょうど目の前で。〔No.386〕"),

("a complete stranger.","見ず知らずの人（完全に知らない人。〔No.387〕"),

("I can hardly wait for the summer.","夏が待ち遠しい（ほとんど待てない）。〔No.388〕"),

("nearly 100 people.","100人近い人々（100は超えない）。〔No.389〕"),

("I am partly to blame.","私にも一部責任がある。〔No.390〕"),

("Jim looks happy, but actually he's sad.","ジムは幸せそうだけど本当は悲しい。〔No.391〕"),

("He is indeed a kind person.","彼は本当に親切な人です。〔No.392〕"),

("even Homer sometimes nods.","《ことわざ》弘法にも筆の誤り（ホメロスでさえも時には居眠りする）。〔No.393〕"),

("Where exactly did you stay in Australia?","オーストラリアの正確にはどこに滞在したのですか。〔No.394〕"),

("Prices are gradually increasing.","物価が徐々に増加している。〔No.395〕"),

("Even experts are only human, and therefore mistakes sometimes occur.","専門家でもただの人間だ．それゆえ，ミスが時々起こる。〔No.396〕"),

("If you can't make it, your mother may go instead.","もしあなたが行けないのなら，君のお母さんが代わりに行ってもよい。〔No.397〕"),

("We are open until 11 p.m. on weekends.","平日は午後11時まで営業しています。〔No.398〕"),

("learn another foreign language besides English.","英語以外の（英語に加えて）外国語を勉強する。〔No.399〕"),

("Everybody was asleep except (for) him.","彼以外は皆眠っていた。〔No.400〕"),

("the debate on whaling.","捕鯨に関する討論。〔No.401〕"),

("Stop criticizing my cooking!","私の料理を批判するのをやめて。〔No.402〕"),

("accuse her of lying.","嘘をついたことで彼女を非難する。〔No.403〕"),

("insist on returning home.","家に帰ると言い張る。〔No.404〕"),

("object to my studying abroad.","私の留学に反対する。〔No.405〕"),

("protest against nuclear tests.","核実験に抗議する。〔No.406〕"),

("highly controversial issue.","非常に論争になっている問題 。〔No.407〕"),

("He is bound to be in the classroom.","今彼はきっと教室にいるはずだ。〔No.408〕"),

("I'll bet (that) he's coming.","きっと彼は来ると思う。〔No.409〕"),

("congratulate her on her exam results.","試験結果について彼女を祝福する。〔No.410〕"),

("praise the man for his diligence.","その男の勤勉さを（その男性を勤勉さで）褒める。〔No.411〕"),

("It's an honor to be here.","お招き頂き有り難うございます.（ここにいることは名誉です）。〔No.412〕"),

("chat over the tea.","お茶を飲みながらおしゃべりをする。〔No.413〕"),

("refer to my childhood.","私の子供の頃について触れる。〔No.414〕"),

("You'd better not mention it.","それには触れないほうがいい（それについて述べないほうがいい）。〔No.415〕"),

("convey my feelings by touching her hand.","彼女の手に触れて, 気持ちを伝える。〔No.416〕"),

("emphasize the importance of breakfast.","朝食の重要性を強調する。〔No.417〕"),

("greatly exaggerate a story.","話をとても大げさに言う。〔No.418〕"),

("I asked him why, but he didn't reply.","彼になぜかと尋ねたが, 彼は返事をしなかった。〔No.419〕"),

("respond to a question.","質問に回答する。〔No.420〕"),

("whisper in her ear.","彼女の耳元でささやく。〔No.421〕"),

("his casual remark.","彼の何気なく言った発言。〔No.422〕"),

("observe the behavior of the birds.","鳥の行動を観察する。〔No.423〕"),

("Your plan is excellent in theory.","あなたの計画は理論上は素晴らしい.。〔No.424〕"),

("handwriting analysis.","筆跡鑑定。〔No.425〕"),

("carry out a chemical experiment .","化学の実験をする。〔No.426〕"),

("a drone with a GPS device.","GPS装置を搭載したドローン。〔No.427〕"),

("natural phenomena.","自然現象。〔No.428〕"),

("a cancer-causing substance.","発がん性物質。〔No.429〕"),

("chemicals that are harmful to the environment.","環境に有害な化学物質。〔No.430〕"),

("burn fossil fuels.","化学燃料を燃やす 。〔No.431〕"),

("a nuclear power plant.","原子力発電所。〔No.432〕"),

("statistics show that younger people prefer football to baseball.","統計によると, 若年層は野球よりサッカーを好むそうだ。〔No.433〕"),

("pursue a career in journalism.","報道関係の道に進む。〔No.434〕"),

("accomplish this task.","この仕事をやり遂げる。〔No.435〕"),

("overcome my shyness.","私の内気なところを克服する。〔No.436〕"),

("fulfill a childhood dream.","子供の頃の夢を果たす 。〔No.437〕"),

("devote my life to helping others.","人助けに人生を捧げる。〔No.438〕"),

("aim to finish by Wednesday.","水曜日までに終わらせるつもりだ。〔No.439〕"),

("face a big challenge.","大きな課題に直面する。〔No.440〕"),

("by trial and error.","試行錯誤で。〔No.441〕"),

("I felt dizzy when I stood up.","立ち上がった時にめまいがした。〔No.442〕"),

("He suddenly went pale.","彼は突然顔面蒼白になった。〔No.443〕"),

("have a hearty appetite.","食欲が旺盛だ。〔No.444〕"),

("I'm starving. What's for dinner?","お腹がぺこぺこだ。夕食は何？。〔No.445〕"),

("mental and physical health.","心身の健康。〔No.446〕"),

("take a rest.","(ちょっと)休憩する 。〔No.447〕"),

("have a stomach ache.","お腹が痛い(お腹の痛みを持つ)。〔No.448〕"),

("have heart bypass surgery.","心臓バイパス手術を受ける。〔No.449〕"),

("the risk of heart disease.","心臓病のリスク。〔No.450〕"),

("the symptoms of heatstroke.","熱中症の症状。〔No.451〕"),

("get cancer.","癌になる。〔No.452〕"),

("fire engines and ambulances.","消防車と救急車。〔No.453〕"),

("recover from a heart attack.","心臓発作から回復する。〔No.454〕"),

("have a sore throat from a cold.","かぜでのどが痛い。〔No.455〕"),

("My eyelids swelled up.","まぶたが腫れ上がった。〔No.456〕"),

("cough loudly.","大きく咳をする。〔No.457〕"),

("My nose has started to bleed.","鼻血が出てきた（鼻が出血し始めた）。〔No.458〕"),

("faint pencil lines.","かすかな鉛筆の跡。〔No.459〕"),

("I am exhausted from walking all day.","1日中歩いてくたくただ。〔No.460〕"),

("cure you of your disease.","あなたの病気を治す。〔No.461〕"),

("He is physically disabled.","彼は身体に障がいがある 。〔No.462〕"),

("Desk work gives me stiff back.","デスクワークで腰が痛い。〔No.463〕"),

("develop my arm muscles.","腕の筋肉をつける。〔No.464〕"),

("stick my tongue out at him.","彼にあかんべえをする（舌を出す。〔No.465〕"),

("lose my sense of time.","時間の感覚をなくす。〔No.466〕"),

("wipe the sweat off my forehead.","額の汗をぬぐう。〔No.467〕"),

("people of the same nationality.","同じ国籍の人々。〔No.468〕"),

("the population of senior citizens.","高齢者（高齢の市民）人口。〔No.469〕"),

("civil rights.","公民権（選挙権などの権利）。〔No.470〕"),

("racial discrimination.","人種差別。〔No.471〕"),

("a domestic flight.","（飛行機の）国内線 。〔No.472〕"),

("show an interest in a rural lifestyle.","田舎暮らしに興味を示す。〔No.473〕"),

("live in the suburbs of Yokohama City.","横浜市の郊外に住む。〔No.474〕"),

("cross the border.","国境を越える。〔No.475〕"),

("a financial burden.","財政的な負担。〔No.476〕"),

("have a negative impact on the tourist industry.","観光業にマイナスの影響を与える。〔No.477〕"),

("women's social status.","女性の社会的地位 。〔No.478〕"),

("All people are created equal.","すべての人は平等に作られている。〔No.479〕"),

("have good relationships with other people.","他者との良好な人間関係をもつ。〔No.480〕"),

("have a good reputation for its beautiful night view.","夜景がきれいなことで評判がよい。〔No.481〕"),

("a trend toward hiring cheap labor.","安価な労働者を雇う風潮。〔No.482〕"),

("public services.","公共事業。〔No.483〕"),

("have a religion.","宗教を信じている（宗教をもっている）。〔No.484〕"),

("a moral responsibility.","道徳的責任。〔No.485〕"),

("our standard of living.","私たちの生活水準。〔No.486〕"),

("the future prosperity of our country.","我が国の未来の繁栄。〔No.487〕"),

("face a serious economic crisis.","深刻な経済危機に直面する。〔No.488〕"),

("LGBT people face prejudice in this workplace.","この職場ではLGBTの人々は偏見の目で見られている。〔No.489〕"),

("fight racial discrimination.","人種差別と戦う。〔No.490〕"),

("do a lot of work for charity.","多くの慈善活動を行う。〔No.491〕"),

("Tourism has brought great benefits to this village.","観光がこの村に多大な恩恵をもたらした。〔No.492〕"),

("work for social welfare.","社会福祉のために働く。〔No.493〕"),

("the Polish community in the U.K.","イギリスのポーランド人社会。〔No.494〕"),

("the rights of the individual.","個人の権利。〔No.495〕"),

("a government official.","政府の役人。〔No.496〕"),

("immigrants from India.","インドからの移民。〔No.497〕"),

("take part in volunteer activities.","ボランティア活動に参加する。〔No.498〕"),

("contribute to world peace.","世界平和に貢献する。〔No.499〕"),

("abolish capital punishment.","死刑制度を廃止する。〔No.500〕"),

("The government should impose a total ban on cigarette advertising.","政府はタバコの広告を全面禁止すべきだ（禁止を課すべきだ）。〔No.501〕"),

("have free access to a computer.","コンピュータを自由に利用できる（利用する権利を持つ）。〔No.502〕"),

("fulfill my duti.","義務を果たす。〔No.503〕"),

("We have nine years of compulsory education in Japan.","日本では義務教育は9年間だ。〔No.504〕"),

("I am responsible for what happens at school.","私には学校で起きることに対して責任がある。〔No.505〕"),

("China's one-child policy.","中国の一人っ子政策 。〔No.506〕"),

(" I was elected a class representative.","私はクラス委員に選ばれた。〔No.507〕"),

("commerce and industry.","商工業。〔No.508〕"),

("my monthly income.","私の月収。〔No.509〕"),

("The restaurant's daily profit is about $1,000.","そのレストランの1日の利益はおよそ千ドルだ。〔No.510〕"),

("pay taxes on my income.","収入に課された税金を払う 。〔No.511〕"),

("This money will cover all your expenses.","このお金があなたのすべての費用を賄うだろう。〔No.512〕"),

("pay back a debt of 100 dollars.","100ドルの借金を返す。〔No.513〕"),

("have a large deposit in the bank.","銀行に多額の預金がある。〔No.514〕"),

("Goods are delivered free of charge.","商品は無料でお届けします。〔No.515〕"),

("get a daily wage of $100.","日給100ドルを得る。〔No.516〕"),

("close down my inn because of the continuing recession.","長引く不況のため旅館を閉める。〔No.517〕"),

("consume a large amount of electricity.","大量の電気を消費する。〔No.518〕"),

("Don't waste so much time on video games.","テレビゲームにそんなに多くの時間を浪費してはいけません。〔No.519〕"),

("invest one million yen in stocks.","株に100万円を投資する。〔No.520〕"),

("import beef from the U.S.","アメリカから牛肉を輸入する。〔No.521〕"),

("receive financial support.","経済的支援を受ける 。〔No.522〕"),

("hire movers to do all the work.","引越し業者を雇って全部任せる。〔No.523〕"),

("employ minors.","未成年を雇う。〔No.524〕"),

("resign as coach.","コーチを辞める。〔No.525〕"),

("qualify for bank loans.","銀行ローンを受ける資格がある。〔No.526〕"),

("assign that important job to him.","彼にその大切な仕事を割り当てる。〔No.527〕"),

("my name, address, and occupation.","氏名、住所、職業。〔No.528〕"),

("start a career as a doctor.","医者として働き始める（職業を始める）。〔No.529〕"),

("enter the legal profession.","法律関係の仕事に就く。〔No.530〕"),

("look over the documents.","資料に目を通す。〔No.532〕"),

("the toy (department).","おもちゃ売り場。〔No.533〕"),

("the branch of this bank.","この銀行の支店。〔No.534〕"),

(" a recommended retail price.","希望小売り価格。〔No.535〕"),

("the nervous system.","神経組織。〔No.536〕"),

("a wooden structure.","木造建造物。〔No.537〕"),

("study architecture.","建築（様式）を勉強する。〔No.538〕"),

("The new city hall is under construction.","新しい市役所が建築中です。〔No.539〕"),

("bodily functions.","身体の機能。〔No.540〕"),

("the surface of the moon.","月面（月の表面）。〔No.541〕"),

("learn about various aspects of Japanese culture.","さまざまな日本文化（日本文化のさまざまな側面）を学ぶ 。〔No.542〕"),

("sit on the edge of the bed.","ベッドの端に座る。〔No.543〕"),

(" The audience consisted mainly of the young people.","観客は大半が若者だった（若者で構成されていた）。〔No.544〕"),

("Water is composed of hydrogen and oxygen.","水は水素と酸素から構成される。〔No.545〕"),

("attach a file to email.","Eメールにファイルを添付する 。〔No.546〕"),

("connect the printer to my PC.","パソコンにプリンターを接続する。〔No.547〕"),

("These two events are closely related.","これら2つの出来事は密接に関連している。〔No.548〕"),

("associate brand names with high quality.","ブランド名から高品質を連想する。〔No.549〕"),

("stick posters on the wall.","stick。〔No.550〕"),

("separate fish from the bones.","魚の骨と身を分ける。〔No.551〕"),

("A thief broke into the building.","泥棒がその建物に侵入した。〔No.552〕"),

("the number of crimes in Japan.","日本の犯罪件数 。〔No.553〕"),

("his motive for the crime.","彼の犯行の動機。〔No.554〕"),

("severely punish him for breaking the rules.","規則を破ったことに対して彼を厳しく罰する。〔No.555〕"),

("violate international law.","国際法を破る 。〔No.556〕"),

("In Japan, the current legal drinking age is 20.","日本では, 現在, 飲酒が許されているのは20歳から（合法の飲酒ができる年齢は20歳）です。〔No.557〕"),

("fight with the enemy.","その敵と戦う。〔No.558〕"),

("compete with each other for good grades.","良い成績を目指してお互いに競い合う。〔No.559〕"),

(" Our team was completely defeated.","私たちのチームは完敗した（完全に打ち負かされた）。〔No.560〕"),

("Over 30 people fell victim to the terrorist attack.","30名以上の人がそのテロの犠牲になった。〔No.561〕"),

("remove an obstacle to an agreement.","合意への障害を取り除く。〔No.562〕"),

("Smoking does you harm.","喫煙は害を及ぼす。〔No.563〕"),

("invade our privacy.","私たちのプライバシーを侵害する。〔No.564〕"),

("endanger the lives of the passengers.","乗客の生命を危険にさらす。〔No.565〕"),

("The meeting was temporarily interrupted by a blackout.","。〔No.566〕"),

("Ben's sudden arrival spoiled our plans.","ベンが突然来たので私たちの計画が台無しになった（私たちの計画を台無しにした）。〔No.567〕"),

("The football match attracted more than 10,000 spectators.","そのサッカーの試合には1万人以上の観客が押し寄せた。〔No.568〕"),

("one of my relatives.","親戚の1人。〔No.569〕"),

("enter the building through the front door.","正面からその建物に入る 。〔No.570〕"),

("accompany my boss to Germany.","上司に同行してドイツに行く。〔No.571〕"),

("the flight's departure.","飛行機の出発。〔No.572〕"),

("get to my destination.","目的地に到着する。〔No.573〕"),

("by public transportation.","。〔No.574〕"),

("fly to Zurich via Hong Kong.","香港経由でチューリッヒに飛ぶ。〔No.575〕"),

("Laziness leads to failure.","怠けると失敗する（失敗に至る）。〔No.576〕"),

("The sun rises in the east.","太陽は東から昇る。〔No.577〕"),

("lower the voting age from 20 to 18.","選挙権年齢を20歳から18歳に引き下げる 。〔No.578〕"),

("This river flows into the Pacific.","この川は太平洋に流れ込む。〔No.579〕"),

("Japan'e economic bubble burst in the 1990s.","日本のバブル経済は1990年代にはじけた。〔No.580〕"),

("The wall cracked in several places.","壁はいくつかの場所でひび割れていた。〔No.581〕"),

("The ice cream quickly melted.","アイスクリームはすぐに溶けてしまった。〔No.582〕"),

("make great progress in English.","英語が随分と上達する。〔No.583〕"),

("advances in technology.","科学技術の進歩。〔No.584〕"),

("deal with the problem.","その問題を扱う。〔No.585〕"),

("handle stress well.","ストレスにうまく対処する。〔No.586〕"),

("cope with unexpected situations.","予期せぬ事態にうまく対処する。〔No.587〕"),

("treat children fairly.","公平に子供を扱う。〔No.588〕"),

("find a clue to his whereabouts.","彼の居所の手がかりを見つける。〔No.589〕"),

("restrict the amount of carbohydrate.","炭水化物の量を制限する。〔No.590〕"),

("limit class size to fifteen.","クラスの人数を15名に制限する。〔No.591〕"),

("forbid employees to accept tips from customers.","従業員が客からのチップを受け取ることを禁じる 。〔No.592〕"),

("impose a total ban on smoking in public places.","公共の場所での喫煙を全面的に禁止する 。〔No.593〕"),

("flatly refuse the offer.","きっぱりとその申し出を断る。〔No.594〕"),

("reject the idea.","その考えを拒絶する。〔No.595〕"),

(" It is no use trying to persuade him to eat carrots.","ニンジンを食べるよう彼を説得するのは無駄だ。〔No.596〕"),

("convince him that I am right.","私が正しいことを彼に確信させる。〔No.597〕"),

("His remark inspired me to study.","彼の発言で私の勉強のスイッチが入った（彼の発言が私を奮起させた）。〔No.598〕"),

("discourage her from buying expensive clothes.","彼女に高い服を買うのを思いとどまらせる。〔No.599〕"),

("promote healthy eating habits.","健全な食生活を促進する 。〔No.600〕"),

("Water expands when it freezes.","水は凍ると膨張する。〔No.601〕"),
       
("extend the deadline by one week.","締め切りを１週間延ばす。〔No.602〕"),

("Travel broadens your mind.","旅は視野を広げてくれる。〔No.603〕"),

("spread a handkerchief over my lap.","ひざの上にハンカチを広げる。〔No.604〕"),

("tie up old magazines with string.","ひもで古雑誌を縛る。〔No.605〕"),

("bind two communities together.","つの社会を結びつける。〔No.606〕"),

("fasten my seat belt.","シートベルトを締める。〔No.607〕"),

("fix the camera to the tripod.","カメラを三脚に固定する。〔No.608〕"),

("install a vending machine in the school cafeteria.","学生食堂に自動販売機を設置する。〔No.609〕"),

("cannot resist buying new shoes.","ついつい新しい靴を買ってしまう。〔No.610〕"),

("obey my parents.","親の言うことに従う。〔No.611〕"),

(" engage in relief operations.","救援活動に従事する。〔No.612〕"),

("bump into a pile of books.","本の山にぶつかる。〔No.613〕"),

("crash into the barrier.","ガードレールに激突する。〔No.614〕"),

("bend down to pick up the pen.","ペンを拾い上げるために腰をかがめる。〔No.615〕"),

("The players hugged each other tightly.","しっかり抱き合った。〔No.616〕"),

("stare at a computer screen.","コンピュータの画面をじっと見つめる 。〔No.617〕"),

(" gaze at the ceiling.","天井を見つめる。〔No.618〕"),

("glance at the clock.","時計をちらっと見る。〔No.619〕"),

("glimpse her face.","彼女の顔がちらりと見える。〔No.620〕"),

(" stretch my arms.","腕を伸ばす。〔No.621〕"),

("stumble  over a rock.","石につまづく。〔No.622〕"),

("press the button.","ボタンを押す。〔No.623〕"),

("drag the table into the kitchen.","台所までそのテーブルを引きずる。〔No.624〕"),

("lean against the wall.","壁にもたれる。〔No.625〕"),

("scratch my back.","背中をひっかく。〔No.626〕"),

("bow to each other.","お互いにおじぎをする。〔No.627〕"),

("She nodded and smiled.","彼女はうなずき、微笑んだ。〔No.628〕"),

('He sighed saying, "You win."',"彼はため息をついた。〔No.629〕"),

("stretch and yawn loudly.","な声であくびをする。〔No.630〕"),

("bury a time capsule at the foot of the cherry tree.","桜の木の根元にムカプセルを埋める。〔No.631〕"),

(" perform difficult tasks.","困難な仕事を遂行する。〔No.632〕"),

("adopt his plan.","彼の計画を採用する。〔No.633〕"),

("escape from the burning house.","燃えさかる家から逃げる。〔No.634〕"),

(" His clothes were scattered all over the floor.","彼の服が床中に脱ぎ散らかされていた。〔No.635〕"),

("fold the paper along the dotted line.","点線に添って紙を折る。〔No.636〕"),

("hang the washing on the pole.","物干しざおに洗濯物を掛ける。〔No.637〕"),

("release the hostages.","人質を解放する。〔No.638〕"),

("strike him on the cheek.","彼のほほを打つ。〔No.639〕"),

("beat a drum.","太鼓をたたく。〔No.640〕"),

("protect our skin from the sun.","日差しから肌を守る。〔No.641〕"),

("twist a wire.","針金をねじ曲げる。〔No.642〕"),

("stir my coffee with a spoon.","スプーンでコーヒーを混ぜる。〔No.643〕"),

("Don't shake the bottle of soda water.","炭酸の瓶を振るな。〔No.644〕"),

("Kinkakuji Temple has burned down several times.","金閣寺は何度か焼け落ちている。〔No.645〕"),

("skip cram school.","塾をサボる。〔No.646〕"),
("Children should be exposed to different culture.","子どもはさまざまな文化に触れる（さらされる）べきだ。〔No.647〕"),
("dip sashimi in soy sauce.","刺身をしょう油につける。〔No.648〕"),
("polish my glasses with a piece of cloth.","布で眼鏡（のレンズ）を磨く。〔No.649〕"),
("cheat in an exam.","試験でカンニングをする。〔No.650〕"),
("attend my sister's wedding.","姉の結婚式に出席する。〔No.651〕"),
("participate in the National Sports Festival.","国体（国民体育大会）に参加する。〔No.652〕"),
("imitate the American teacher's English.","そのアメリカ人教師の英語をまねる。〔No.653〕"),
("believe that Santa Claus exists.","サンタクロースは存在すると信じる。〔No.654〕"),
("Several problems arose.","いくつかの問題が生じた。〔No.655〕"),
("The accident occurred at this intersection at about 10 p.m.","その事故は午後10時ごろにこの交差点で起きた。〔No.656〕"),
("generate electricity.","電気を生み出す。〔No.657〕"),
("I was involved in a traffic accidennt.","交通事故に巻き込まれた。〔No.658〕"),
("This job requires many year's experience.","この仕事には長年の経験が必要だ。〔No.659〕"),
("The Japanese officials discussed the issue with ther French counterparts.","日本政府高官はフランス政府の高官（対応する人々）とその問題について話し合った。〔No.660〕"),
("maintain high standards.","高水準を維持する。〔No.661〕"),
("The meeting lasted (for) three hours.","会議は3時間続いた。〔No.662〕"),
("If you fever persists, you should see a doctor.","熱が続くようなら、医者に行ったほうがいい。〔No.663〕"),
("I have gained five kilos.","5キロ太った。〔No.664〕"),
("obtain a work visa.","就労ビザを得る。〔No.665〕"),
("acquire a foreigh language.","外国語を習得する。〔No.666〕"),
("examine the old records.","古い記録を調べる。〔No.667〕"),
("search the house for my earrings.","イヤリングがないか家を捜す。〔No.668〕"),
("a high school class reunion.","高校の同窓会。〔No.669〕"),
("study abroad on a scholarship.","奨学金で留学する。〔No.670〕"),
("There is no logic in your argument.","君の主張には論理がない。〔No.671〕"),
("follow the on-screen instructions.","画面上の指示に従う。〔No.672〕"),
("People's lifestyles are often determined by their incomes.","人々の生活スタイルは収入で決まることが多い。〔No.673〕"),
("conclude that the factory should be closed.","その工場は閉鎖すべきだと結論を下す。〔No.674〕"),
("distinguish sheep from goats.","ヒツジをヤギと区別する。〔No.675〕"),
("classify the books according to subject [their subject(s)].","テーマで図書を分類する。〔No.676〕"),
("It is estimated that this shrine is over 500 years old.","この神社の歴史は500年以上だと見積もられている。〔No.677〕"),
("Fred was selected for the national team.","フレッドはナショナルチームに選ばれた。〔No.678〕"),
("organize my thoughts before speaking.","話す前に私の考えをまとめる。〔No.679〕"),
("Do you recognize me?","私が誰だかわかりますか。〔No.680〕"),
("You are supposed to take off your shoes at the door.","玄関では靴を脱ぐことになっています。〔No.681〕"),
("assume that all people are the same.","人間は皆同じだと思い込む。〔No.682〕"),
("I don't care about that at all.","私はそれをまったく気にしません。〔No.683〕"),
("I have never once doubted him.","彼のことを疑ったことは一度もない。〔No.684〕"),
("I noticed that there was a hole in my sock.","靴下に穴が空いていることに気がついた（＊その後も気がついていた〈状態〉）。〔No.685〕"),
("I am aware that my time on earth is limited.","この世の時間には限りがあると気づいている。〔No.686〕"),
("I was conscious of the fact that it was an important meeting.","それが重要な会議であることは意識していた。〔No.687〕"),
("The professor is concerned with environmental problems.","その教授は環境問題に関心を持っている。〔No.688〕"),
("Zack was born in the U.S. but he regards Japan as his home.","ザックはアメリカ生まれだが、日本を故郷だと思っている。〔No.689〕"),
("We are fully committed to the project.","私たちはその計画に専念している。〔No.690〕"),
("memorize the times tables.","九九を暗記する。〔No.691〕"),
("Her parents approved of her marriage.","彼女の両親は彼女の結婚を認めた。〔No.692〕"),
("forgive him for what he said.","彼の発言に対して彼を許す。〔No.693〕"),
("take cars for granted.","車の存在を当然（認められたもの）と考える。〔No.694〕"),
("I recalled that I had seen him cheating on the test.","彼がテストでカンニングするのを目撃したことを思い出した。〔No.695〕"),
("abandon my dream of studying abroad.","留学する夢を諦める。〔No.696〕"),
("get rid of my old toys.","古いおもちゃを処分する。〔No.697〕"),
("eliminate sex discrimination.","男女差別をなくす。〔No.698〕"),
("relieve my stress.","ストレスを発散させる。〔No.699〕"),
("Illegally parked bicycles will be removed.","放置自転車（不法駐輪された自転車）は撤去される。〔No.700〕"),
("develop natural resources.","天然資源を開発する。〔No.701〕"),
("conservation groups.","環境保護団体。〔No.702〕"),
("preserve endangered species.","絶滅危惧種を保護する。〔No.703〕"),
("natural disasters in Japan.","日本の自然災害。〔No.704〕"),
("creatures from another planet.","別の惑星から来た生き物。〔No.705〕"),
("the natural environment.","自然環境。〔No.706〕"),
("The sun appeared on the horizon.","太陽が水平線上に現れた。〔No.707〕"),
("organic agriculture.","有機農業。〔No.708〕"),
("the main crops in this area.","この地域の主要な農産物。〔No.709〕"),
("cultivate the soil.","土を耕す。〔No.710〕"),
("pull up weeds.","雑草を抜く。〔No.711〕"),
("he almost drowned, but luckily he was saved.","彼は溺れかけたが、幸い救助された。〔No.712〕"),
("This roof leaks.","雨漏りがする（この屋根は濡れる）。〔No.713〕"),
("This town has a mild[harsh] climate.","この町の気候は温暖[過酷]です。〔No.714〕"),
("ice crystals in the atmosphere.","大気中の氷の結晶。〔No.715〕"),
("according to the weather forecast.","天気予報によれば。〔No.716〕"),
("The thermometer read ten degrees below zero.","温度計が氷点下10度を指した。〔No.717〕"),
("Kyoto is hot and humid in summer.","京都は夏は蒸し暑い。〔No.718〕"),
("tropical rainforests.","熱帯雨林。〔No.719〕"),
("install solar panels on the roof.","屋根にソーラーパネルを取り付ける。〔No.720〕"),
("I was bitten by a mosquito.","蚊に刺された。〔No.721〕"),
("elephants and other endangered species.","ゾウなどの絶滅危惧種。〔No.722〕"),
("Dinosaurs became extinct millions of years ago.","恐竜は何百万年も前に絶滅した。〔No.723〕"),
("feed these goldfish.","これらの金魚にえさを与える。〔No.724〕"),
("energetic boys and girls.","活発な少年少女たち。〔No.725〕"),
("A greedy child ate all the pies.","欲張りな子がパイを全部食べた。〔No.726〕"),
("a brave firefighter.","勇敢な消防士。〔No.727〕"),
("leave a generous tip.","気前よく（気前のよい額の）チップを置く。〔No.728〕"),
("intellectual property.","知的財産。〔No.729〕"),
("Babies are curious about everything.","赤ん坊はすべての物に好奇心が強い。〔No.730〕"),
("Children are very imaginative.","子どもは非常に想像力が豊かだ。〔No.731〕"),
("a talkative taxi driver.","おしゃべりなタクシー運転手。〔No.732〕"),
("I cannot afford to buy a new car.","新車を買う余裕がない。〔No.733〕"),
("Beth is eager to buy a fur coat.","ベスは毛皮のコートを買いたがっている。〔No.734〕"),
("Don't be selfish.","自分勝手なことをしてはだめだよ。〔No.735〕"),
("Her attitude suddenly became aggressive.","彼女の態度が突然攻撃的になった。〔No.736〕"),
("Never be cruel to animals.","動物を残酷に扱うな。〔No.737〕"),
("Lucy is addicted to her smartphone.","ルーシーはスマホ（スマートフォン）中毒である。〔No.738〕"),
("a stubborn old man.","頑固な老人。〔No.739〕"),
("make an earnest effort.","真面目に努力する。〔No.740〕"),
("a dress with a bold design.","大胆なデザインの服。〔No.741〕"),
("feel guilty about lying to her.","彼女にうそをついたことを申し訳なく思う。〔No.742〕"),
("Jack is innocent of the crime.","ジャックはその犯罪に関して無実だ。〔No.743〕"),
("make a sincere effort.","ひたむきな（誠実な）努力をする。〔No.744〕"),
("Paul is modest about his success.","ポールは成功を鼻にかけない。〔No.745〕"),
("make a stupid mistake.","ばかな間違いをする。〔No.746〕"),
("Rick is indifferent to politics.","リックは政治に無関心だ。〔No.747〕"),
("Japanese trains are extremely punctual.","日本の列車は非常に時間に正確だ。〔No.748〕"),
("My son is a coward when it comes to going to the dentist.","息子は歯医者に行くことにかけては臆病者です。〔No.749〕"),
("Each plane has to follow a precise route.","どの飛行機も正確な航路をたどらなくてはならない。〔No.750〕"),
("an accurate map.","正確な地図。〔No.751〕"),
("the proper use of chopsticks.","箸の正しい使い方。〔No.752〕"),
("keep the office tidy.","事務所をきちんとしておく。〔No.753〕"),
("neat piles of towels.","きちんと積まれたタオル。〔No.754〕"),
("an efficient method of transporting goods.","商品を輸送する能率的な方法。〔No.755〕"),
("What you're saying is reasonable.","あなたの言っていることはもっともだ。〔No.756〕"),
("a significant change in our plans.","我々の計画の重大な変更。〔No.757〕"),
("Don't waste your precious time.","貴重な時間を浪費するな。〔No.758〕"),
("Water is essential to living things.","水は生物にとって不可欠だ。〔No.759〕"),
("make a fundamental change.","根本的な改革を行う。〔No.760〕"),
("in a critical condition.","危篤状態（危機的な状況）で。〔No.761〕"),
("do serious damage to the environment.","環境に深刻な打撃を与える。〔No.762〕"),
("a highly complex process.","非常に複雑な過程。〔No.763〕"),
("make things complicated[×complex].","事態を複雑にする。〔No.764〕"),
("the delicate question of salary.","給与という繊細な問題。〔No.765〕"),
("write in plain English.","平易な英語で書く。〔No.766〕"),
("It was obvious that she was unwell.","彼女の体調がすぐれないのは明白だった。〔No.767〕"),
("make remarkable progress.","注目すべき進歩を遂げる。〔No.768〕"),
("an outstanding shogi player.","傑出した棋士。〔No.769〕"),
("various ideas.","さまざまな考え。〔No.770〕"),
("diverse culture.","多様な文化。〔No.771〕"),
("What sort of soap do you use?.","どのような（種類の）石けんを使っているの。〔No.772〕"),
("a fancy French restaurant.","高級なフレンチレストラン。〔No.773〕"),
("have a marvelous memory.","驚くべき記憶力を持っている。〔No.774〕"),
("a fabulous hotel in Hawaii.","ハワイのとても素敵なホテル。〔No.775〕"),
("lead an active school life.","活動的な学生生活を送る。〔No.776〕"),
("positive  thinking.","前向きな考え。〔No.777〕"),
("a pure wool blanket.","100％（純粋な）ウールの毛布。〔No.778〕"),
("make steady progress.","着実に進歩する（着実な進歩をする）。〔No.779〕"),
("Dave is flexible about everything.","デイブはあらゆることに柔軟に対応する。〔No.780〕"),
("a ripe mango.","熟したマンゴー。〔No.781〕"),
("Ann is mature for her age.","アンは年の割には成熟している。〔No.782〕"),
("His temperature remains stable.","彼の体温は安定している。〔No.783〕"),
("the negative aspects of aging.","歳をとることの否定的な側面。〔No.784〕"),
("a vague description.","曖昧な説明。〔No.785〕"),
("Steve is a little weird, isn't he?.","スティーブってちょっと変だよね。〔No.786〕"),
("a rough road.","でこぼこの道。〔No.787〕"),
("severe criticism.","厳しい評判。〔No.788〕"),
("He is quite passive, and never speaks out.","彼はかなり消極的で、決して意見をはっきり言わない。〔No.789〕"),
("His effort was in vain.","彼の努力は無駄になった。〔No.790〕"),
("a fake diamond.","偽物のダイヤモンド。〔No.791〕"),
("It is risky to buy a used car from a private seller.","個人の売り手から中古車を買うのは危険だ。〔No.792〕"),
("an absurd idea.","ばかげた考え。〔No.793〕"),
("an odd habit.","奇妙な習慣。〔No.794〕"),
("the contrast between the two.","その両者の対比。〔No.795〕"),
("a valid reason.","妥当な理由。〔No.796〕"),
("This T-shirt is available in all sizes.","このTシャツはすべてのサイズでご用意できます（手に入る）。〔No.797〕"),
("feel more relaxed in casual clothes.","ふだん着（気楽な服）のほうが落ち着く。〔No.798〕"),
("collect rate stamps.","珍しい切手を集める。〔No.799〕"),
("gain practical experience.","実践経験を積む。〔No.800〕"),
("brand-new climbing boots.","新品の登山靴。〔No.801〕"),
("rapid economic growth.","急速な経済成長。〔No.802〕"),
("There is an urgent need for more nursery schools.","もっと多くの保育所が緊急に必要だ（もっと多くの保育所に対する緊急の必要性がある）。〔No.803〕"),
("a sharp increase in prices.","急激な物価の上昇。〔No.804〕"),
("The supporters stayed calm.","サポーターたちは落ち着いていた。〔No.805〕"),
("a shallow bathtub.","浅い浴槽。〔No.806〕"),
("Her arms were bare.","彼女の腕はむき出しだった。〔No.807〕"),
("swim naked in the river.","川で裸で（裸の状態で）泳ぐ。〔No.808〕"),
("I am independent of my parents.","私は両親の世話になっていない（両親から独立している）。〔No.809〕"),
("a tense atmosphere in the waiting room.","待合室の張りつめた雰囲気。〔No.810〕"),
("have broad shoulders.","肩幅が広い。〔No.811〕"),
("a narrow mountain path.","狭い山道。〔No.812〕"),
("The seat next to mine was vacant.","私の隣の席は空いていた。〔No.813〕"),
("have vivid memories of my time in Sydney.","シドニーで過ごした日々の鮮明な思い出が残っている。〔No.814〕"),
("promote mutual understanding.","相互の理解を促進させる。〔No.815〕"),
("the awkward movements of the robot.","そのロボットのぎこちない動き。〔No.816〕"),
("Historians try to be objective.","歴史家は客観的になるように努める。〔No.817〕"),
("manual work.","手仕事。〔No.818〕"),
("The twins are very alike.","その双子はとてもよく似ている。〔No.819〕"),
("Women tend to count calories more than men do.","女性は男性よりカロリー計算をする傾向がある。〔No.820〕"),
("Ken has trained so hard that he deserves to win.","ケンは一生懸命トレーニングしたので、勝ってもおかしくない（勝つことに値する）。〔No.821〕"),
("glass fragments.","ガラスの破片。〔No.822〕"),
("pass the exam with ease.","その試験に容易に受かる。〔No.823〕"),
("a wide range of topics.","幅広い話題。〔No.824〕"),
("do business on a large scale.","手広く（大規模に）商売をする。〔No.825〕"),
("This custom is unique to Japan.","この習慣は日本特有のものだ。〔No.826〕"),
("a particular situation.","ある特定の状況。〔No.827〕"),
("the general public.","一般大衆。〔No.828〕"),
("the lives of ordinary people.","庶民（ふつうの人々）の生活。〔No.829〕"),
("a typical Japanese-style breakfast.","典型的な日本の朝食。〔No.830〕"),
("Middle-aged and older people account for 70 % of the group.","中高年がその集団の70%を占める。〔No.831〕"),
("calculate this month's expenses.","今月の出費を計算する。〔No.832〕"),
("Please measure your blood pressure.","血圧を測ってください。〔No.833〕"),
("Let's split the bill.","割り勘にしよう（勘定を割る）。〔No.834〕"),
("the volume of a cylinder.","円柱の体積。〔No.835〕"),
("The proportion of boys to girls in my school is five to one.","私の学校の男女比は5：1だ。〔No.836〕"),
("two dozen eggs.","2ダースの卵。〔No.837〕"),
("a large amount of fat.","大量の脂肪。〔No.838〕"),
("a huge mass of data.","非常に多くのデータ。〔No.839〕"),
("There is a water shortage in this area.","この地域は水が不足している。〔No.840〕"),
("an enormous amount of time.","莫大な（量の）時間。〔No.841〕"),
("in my spare time.","余分な時間で。〔No.842〕"),
("arrange for someone to drive her home.","誰かが彼女を家まで車で送ってくれるよう手配する。〔No.843〕"),
("adapt to a new school life.","新しい学校生活に適応する。〔No.844〕"),
("This sweater matches your skirt.","このセーターは君のスカートによく合っている。〔No.845〕"),
("This dress fits you.","このワンピースは（サイズが）君にぴったりだ。〔No.846〕"),
("An emergency has arisen.","緊急事態が生じた。〔No.847〕"),
("on this occasion.","このような場合には。〔No.848〕"),
("The circumstances are changing minute by minute.","状況は刻一刻と変化している。〔No.849〕"),
("The incident occurred at around 3 a.m.","その出来事は午前3時ごろに起きた。〔No.850〕"),
("the current energy crisis.","現在のエネルギー問題。〔No.851〕"),
("a temporary license.","仮免許。〔No.852〕"),
("a permanent member of the U.N. Security Council.","国連安全保障理事会の常任理事国。〔No.853〕"),
("the previous morning.","（ある日の）前日の朝（前の朝）。〔No.854〕"),
("the former and the latter.","前者と後者。〔No.855〕"),
("write the annual report.","年次報告書を作成する。〔No.856〕"),
("contemporary literature.","現代文学。〔No.857〕"),
("Have you seen him lately?.","最近彼と会った。〔No.858〕"),
("immediately after breakfast.","朝食後すぐに。〔No.859〕"),
("supply people with drinking water.","人々に飲み水を供給する。〔No.860〕"),
("AI has replaced humans in many fields.","人工知能が多くの分野で人間に取って代わった。〔No.861〕"),
("exchange Japanese yen for U.S. dollars.","日本円を米ドルと交換する。〔No.862〕"),
("substitute honey for sugar.","砂糖の代わりにはちみつを使う。〔No.863〕"),
("submit an application form.","申込用紙を提出する。〔No.864〕"),
("an  alternative to nuclear power.","原子力の代わりになるもの。〔No.865〕"),
("deliver pizzas.","ピザを配達する。〔No.866〕"),
("enclose several photos.","何枚かの写真を同封する。〔No.867〕"),
("put a stamp on the envelope.","封筒に切手を貼る。〔No.868〕"),
("play tricks on Kevin.","ケビンにいたずらをする。〔No.869〕"),
("carry a heavy load on my back.","重い荷物を背負う。〔No.870〕"),
("the contents of her bag.","彼女のかばんの中身。〔No.871〕"),
("Sweeping the street in front of my house is part of my daily routine.","家の前の道を掃くことは日課の一部です。〔No.872〕"),
("a survey of 2,000 households.","2,000世帯の調査。〔No.873〕"),
("high quality goods.","高品質の商品。〔No.874〕"),
("spend a lot of money on luxuries.","ぜいたく品に多額のお金を使う。〔No.875〕"),
("Do you accept credit cards?.","クレジットカードは使えますか。〔No.876〕"),
("a survey using a questionnaire.","アンケート（を用いた）調査。〔No.877〕"),
("make a reservation under the name of Sophie.","ソフィーの名前で予約する。〔No.878〕"),
("receive hearty applause.","心からの拍手をもらう。〔No.879〕"),
("make a fuss about trivial things.","くだらないことで大騒ぎする。〔No.880〕"),
("the reward for the job.","その仕事に対する報酬。〔No.881〕"),
("have a farewell party.","送別会を開く。〔No.882〕"),
("receive a warm reception.","暖かい歓迎を受ける。〔No.883〕"),
("Thank you for your kind hospitality.","親切なおもてなしに感謝します。〔No.884〕"),
("a large portion of roast beef.","大盛りのローストビーフ。〔No.885〕"),
("cook, clean, and do the laundry.","料理、掃除、洗濯をする。〔No.886〕"),
("remove stubborn oil stains.","頑固な油汚れをとる。〔No.887〕"),
("dye my hair brown.","髪を茶色に染める。〔No.888〕"),
("plug a vacuum cleaner into the outlet.","掃除機をコンセントに差し込む。〔No.889〕"),
("I'll just take a little nap.","ちょっと昼寝をします。〔No.890〕"),
("I woke up at six a.m., but did not get up.","午前6時に目が覚めたが、起きなかった。〔No.891〕"),
("a soft drink vending machine.","清涼飲料水の自動販売機。〔No.892〕"),
("a local grocery (store).","地元の食料雑貨店。〔No.893〕"),
("I have an appointment to see the dentist at five.","5時に歯医者の予約を入れている。〔No.894〕"),
("You'd better consult your doctor.","医者に診てもらいなさい。〔No.895〕"),
("Koyasan is registered as a World Heritage Site.","高野山は世界遺産に登録されている。〔No.896〕"),
("subscribe to the magazine.","雑誌を定期購読する。〔No.897〕"),
("Freedom of speech is guaranteed under the Constitution of Japan.","言論の自由は日本国憲法で保証されている。〔No.898〕"),
("wipe the table.","テーブルを拭く。〔No.899〕"),
("sweep the floor with a broom.","ほうきで床を掃く。〔No.900〕"),
("transfer at Rome.","ローマで乗り換える。〔No.901〕"),
("They have divorcerd.","彼らは離婚した。〔No.902〕"),
("A terrible fate awaited them.","恐ろしい運命が彼らを待ち受けていた。〔No.903〕"),
("It was his destiny to save his nation.","国を救うことが彼の運命だった。〔No.904〕"),
("ice crem with a green tea flavor.","抹茶味のアイスクリーム。〔No.905〕"),
("the perfume of roses.","バラの香り。〔No.906〕"),
("the ingredients of lasagne.","ラザニアの材料。〔No.907〕"),
("a tender sirloin steak.","柔らかいサーロインステーキ。〔No.908〕"),
("a bitter experience.","つらい経験。〔No.909〕"),
("prefer an aisle seat to a window seat.","窓側の席より通路側の席のほうが好き。〔No.910〕"),
("keep track of my schedule.","スケジュール管理をする（スケジュールの経過を追う）。〔No.911〕"),
("an election district.","選挙区。〔No.912〕"),
("indoor sports facilities.","室内スポーツ施設。〔No.913〕"),
("food vendors in the basement of a department store.","デパートの地下の食品売り場。〔No.914〕"),
("line up in order of height.","身長順に並ぶ。〔No.915〕"),
("a distant island.","遠い島。〔No.916〕"),
("live in a remote village.","へんぴな村に住んでいる。〔No.917〕"),
("New Mexico is located 2,240 meters above sea level.","ニューメキシコは海抜2,240メートルの位置にある。〔No.918〕"),
("occupy two seats on the train.","電車で2つの席を占有する。〔No.919〕"),
("a msnsion surrounded by red brick walls.","赤れんがの塀に囲まれたお屋敷。〔No.920〕"),
("classical music.","クラシック音楽。〔No.921〕"),
("an advanced civilization.","高度な文明。〔No.922〕"),
("the cultural heritage of Japan.","日本の文化遺産。〔No.923〕"),
("a film script.","映画の台本。〔No.924〕"),
("a Japanese folk table.","日本の民話。〔No.925〕"),
("the Nobel Prize in Literature.","ノーベル文学賞。〔No.926〕"),
("9/11 was a terrible tragedy.","9/11（の同時多発テロ）は本当に悲劇だった。〔No.927〕"),
("learn a lot of poetry by heart.","多くの詩を暗唱する。〔No.928〕"),
("a biography of Helen Keller.","ヘレン・ケラーの伝記。〔No.929〕"),
('the term "degital native".',"『デジタルネイティブ』という言葉。〔No.930〕"),
("As the proverb goes[says], time flies.","ことわざにもあるように、光陰矢の如しだ。〔No.931〕"),
("He spesks fluent Spanish.","彼はスペイン語が流暢だ。〔No.932〕"),
("translate a sentence literally.","文を直訳する（文字どおりに文を訳す）。〔No.933〕"),
("Obesity is defined as a BMI of 30 or above.","肥満はBMI30以上と定義されている。〔No.934〕"),
("interpret his silence as a refusal.","彼の沈黙を拒否と解釈する。〔No.935〕"),
("quote a passage from the Bible.","聖書の一節を引用する。〔No.936〕"),
('The word bonen-kai literally means "a forget-the-year party".',"「忘年会」は文字どおりには「年を忘れる会」を意味する。〔No.937〕"),
("Research indicates that the medicine can be harmful to children.","研究はその薬が子どもに害があるかもしれないことを示している。〔No.938〕"),
("reveal a secret to her.","彼女に秘密を漏らす。〔No.939〕"),
("announce the results of the election.","選挙の結果を発表する。〔No.940〕"),
("The interview with him was broadcast live.","彼のインタビューは生放送された。〔No.941〕"),
("prove (that) it is false.","それが偽物だと証明する。〔No.942〕"),
("advertise the concert.","コンサートを宣伝する。〔No.943〕"),
("seek specialist advice.","専門家の意見を求める。〔No.944〕"),
("entertain an audience.","観客を楽しませる。〔No.945〕"),
("The clown amused us all.","そのピエロは私たちみんなを楽しませた。〔No.946〕"),
("The pop star's marriage attracted media attention.","人気スターの結婚はマスコミの関心を引きつけた。〔No.947〕"),
("I was fascinated by the Sagrada Familia.","サグラダファミリアに魅了された。〔No.948〕"),
("I was absorbed in a book.","私は本に夢中になっていた。〔No.949〕"),
("I am very fond of ramen.","私はラーメンが大好きだ。〔No.950〕"),
("Oh! You scared me!.","もう、びっくりしたじゃない（あなたは私をおびえさせた）。〔No.951〕"),
("I was alarmed to hear a man shout.","男が叫ぶのを聞いて、びっくりした。〔No.952〕"),
("Sam amazed me with his cooking skills.","サムはその料理の腕で私を驚かせた。〔No.953〕"),
("Her habit of biting her nails irritates me.","彼女の爪をかむ癖は私をいらいらさせる。〔No.954〕"),
("What puzzles me is why she can't understand me.","私を当惑させるのは、なぜ彼女は私の言うことが理解できないかだ。〔No.955〕"),
("His reply confused me.","彼の返答は私を困惑させた。〔No.956〕"),
("I was bored with his long speech.","彼の長い話にうんざりした。〔No.957〕"),
("I felt frustrated bacause I could not express myself well.","うまく表現できなくてもどかしい思いをした。〔No.958〕"),
("I was disappointed with the soccer game.","サッカーの試合にがっかりした。〔No.959〕"),
("I was embarrassed about using the wrong fork at dinner.","夕食時、フォークを使い間違えて恥ずかしかった。〔No.960〕"),
("I'm ashamed of hurting her feelings.","彼女の気持ちを傷つけたことを恥じている。〔No.961〕"),
("feel uneasy about living alone.","ひとり暮らしは不安だ。〔No.962〕"),
("Please do not hesitate to contact me if you have any questions.","質問があればどうぞご遠慮なく連絡して（連絡するのをためらわないで）ください。〔No.963〕"),
("I was reluctant to go with him.","彼と一緒に行くのは気が進まなかった。〔No.964〕"),
("The little girl was trembling with fear.","その少女は恐怖で震えていた。〔No.965〕"),
("Boast about how clever my son is.","私の息子がどれほど賢いかを自慢する。〔No.966〕"),
("be jealous of his promotion.","彼の昇進が妬ましい。〔No.967〕"),
("envy him for his wealth.","財産のことで彼を羨ましく思う。〔No.968〕"),
("yell at the children to be quiet.","子どもたちに静かにせよと叱る（大声で言う）。〔No.969〕"),
("respect Gandhi.","ガンジーを尊敬する。〔No.970〕"),
("a deep sense of despair.","深い絶望感。〔No.971〕"),
("show my emotions.","感情を表に出す。〔No.972〕"),
("have deep sympathy for the victims.","犠牲者たちに本当に同情する。〔No.973〕"),
("It's a shame that you did'nt see that movie.","あなたがその映画を見ていないのは残念です。〔No.974〕"),
("achieve my ambition to be a pilot.","パイロットになるという望みを叶える。〔No.975〕"),
("have[gain, lose] confidence.","自信を持つ[得る、失う]。〔No.976〕"),
("have the courage to say no.","断る勇気を持つ。〔No.977〕"),
("completely ignore him.","完全に彼を無視する。〔No.978〕"),
("I owe my success to him.","私の成功は彼のおかげだ。〔No.979〕"),
("overcome my fear of snakes.","ヘビに対する恐怖を克服する。〔No.980〕"),
("a method for reducing stress.","ストレスを減らすやり方。〔No.981〕"),
("receive a great deal of media attention.","メディアにとても注目される。〔No.982〕"),
("a means of communication.","意思疎通の手段。〔No.983〕"),
("You are correct to som extent.","君はある程度正しい。〔No.984〕"),
("It's extremely hot outside.","外はひどく暑い。〔No.985〕"),
("the total cost.","総費用。〔No.986〕"),
("Our customers are largely women..","私たちのお客様は大部分が女性だ。〔No.987〕"),
("get moderate exercise.","適度な運動をする。〔No.988〕"),
("feel somewhat sad.","いくぶん悲しい気分だ。〔No.989〕"),
("a subtle difference in meaning.","かすかな意味の違い。〔No.990〕"),
("It will take two weeks, possibly longer.","2週間，ひょっとしたらそれ以上かかるかもしれない。〔No.991〕"),
("After changing jobs many times, I eventually found a job in a bank.","何度も転職を繰り返したあと，最終的に，銀行に就職しました。〔No.992〕"),
("Rich people are not necessarily happy.","金持ちが必ずしも幸せとは限らない。〔No.993〕"),
("talk frankly with a friend.","友だちと率直に話し合う。〔No.994〕"),
("ironically, my car was stolen right in front of the police station.","皮肉なことに，私の車は警察署の真ん前で盗まれた。〔No.995〕"),
("Smoking is bad for you. moreover, it costs a lot.","喫煙は体に悪い．その上お金がかかる。〔No.996〕"),
("Hurry up; otherwise you'll be late.","急げ！さもないと遅れるぞ。〔No.997〕"),
("regardless of sex, race, or nationality.","性別，民族，国籍とは無関係に。〔No.998〕"),
("This ring is worth over two million dollars.","この指輪は200万ドル以上の価値がある。〔No.999〕"),
("according to a public poll.","世論調査によると。〔No.1000〕"),

("purpose a new program.","新しい計画を提案する。〔No.1001〕"),
("demand an apology.","謝罪を要求する。〔No.1002〕"),
("desire to study in the U.K..","イギリスに留学したいという強い願望を持つ。〔No.1003〕"),
("dismiss his proposal as unrealistic.","非現実的だとして彼の提案を退ける。〔No.1004〕"),
("(God) bless you!.","お大事に.（神があなたを祝福しますように）。〔No.1005〕"),
("his glory days as a college basketball star.","大学バスケットボールのスター選手としての彼の栄光の日々。〔No.1006〕"),
("compliments on his shoes.","彼の靴への褒め言葉。〔No.1007〕"),
("the wedding feast.","結婚式の宴席。〔No.1008〕"),
("He declared that he was innocent..","彼は無実だとはっきりと述べた。。〔No.1009〕"),
("demonstrate my real ability.","実力を発揮する（真の実力を示す）。〔No.1010〕"),
("highlight the issue of global warming.","地球温暖化の問題を強調する。〔No.1011〕"),
("Alex implied that he would resign..","アレックスは辞意をほのめかした。。〔No.1012〕"),
("recite a poem.","詩を暗唱する。〔No.1013〕"),
("the sun's rays.","太陽光線。〔No.1014〕"),
("The workers were exposed to radiation..","その労働者たちは被爆した（放射線にさらされた）。。〔No.1015〕"),
("laboratory experiments.","研究室での実験。〔No.1016〕"),
("Water is made up of oxygen and hydrogen..","水は酸素と水素からできている。。〔No.1017〕"),
("a water molecule.","水分子。〔No.1018〕"),
("a chemical compound.","化合物。〔No.1019〕"),
("nerve tissue.","神経組織。〔No.1020〕"),
("remove the cancerous cells.","がん細胞を除去する。〔No.1021〕"),
("the gene for black hair.","黒い髪の遺伝子。〔No.1022〕"),
("solid fuel.","固形燃料。〔No.1023〕"),
("Jupiter's sixth satellite.","木星の６番目の衛星。〔No.1024〕"),
("the moon's orbit around the earth.","地球を回る月の軌道。〔No.1025〕"),
("launch a Sun probe.","太陽探査機を打ち上げる。〔No.1026〕"),
("make an attempt to break his record.","彼の記録を破ろうと試みる。〔No.1027〕"),
("have a remarkable capacity to learn language.","目立った言語学習能力を有している。〔No.1028〕"),
("The factory is capable of producing 100 cars per hour.","この工場では１時間に100台もの車を生産できる。〔No.1029〕"),
("I have attained my ideal weight.","私は理想とする体重に達した。〔No.1030〕"),
("Make a desperate effort to succeed.","成功するために必死の努力をする。〔No.1031〕"),
("Struggle to bring up my children.","必死になって子どもたちを育てる（育てるのに苦労する）。〔No.1032〕"),
("She dedicated herself to her job.","彼女は仕事に没頭した（仕事に自分自身を捧げた）。〔No.1033〕"),
("Do you feel any pain?","（医者の発言）痛みはありますか。〔No.1034〕"),
("Work under a lot of strain.","過度の負担の下で働く。〔No.1035〕"),
("A remedy for colds.","かぜの治療法。〔No.1036〕"),
("His face was gray with fatigue.","過労のため, 彼の顔を蒼白だった。〔No.1037〕"),
("The problem of obesity.","肥満の問題。〔No.1038〕"),
("Round-the-clock nursing care.","24時間体制の看護。〔No.1039〕"),
("Terminal cancer.","末期がん。〔No.1040〕"),
("Terry is three-month pregnant.","テリーは妊娠3か月だ。〔No.1041〕"),
("Heal people by laying my hands on their bodies.","体に手を当てて人々を治す。〔No.1042〕"),
("Twist my ankle.","足首をひねる。〔No.1043〕"),
("Stick up my thumb.","親指を立てる。〔No.1044〕"),
("I've got a pimple on my forehead.","（私の）おでこににきびができた。〔No.1045〕"),
("Stick out my chin.","あごを突き出す（＊挑戦的態度の示唆）。〔No.1046〕"),
("Have a chest X-ray examination.","胸部のレントゲン検査を受ける。〔No.1047〕"),
("Early detection of breast cancer.","乳がんの早期発見。〔No.1048〕"),
("The heart and the lungs.","（ヒトの）心肺。〔No.1049〕"),
("Wait for an organ transplant.","臓器移植を待つ。〔No.1050〕"),
("Have good [normal, bad] vision.","視力が良い [正常だ. 悪い]。〔No.1051〕"),
("A model of the human skeleton.","ヒトのがい骨の模型。〔No.1052〕"),
("Lose all sensation in my toes.","足の指先の感覚がなくなる。〔No.1053〕"),
("Does the restaurant have a dress code?","そのレストランに服装規定はありますか。〔No.1054〕"),
("The agenda for today's meeting.","本日の会議の議題。〔No.1055〕"),
("Regardless of age or gender.","年齢, 性別とは無関係に。〔No.1056〕"),
("Fight for liberty and equality.","自由と平等のために戦う。〔No.1057〕"),
("A crime against humanity.","人類に対する犯罪。〔No.1058〕"),
("In the history of mankind.","人類の歴史において。〔No.1059〕"),
("An authority on orthodontics.","歯科矯正学の権威。〔No.1060〕"),
("Have a strong sense of justice.","正義感が強い。〔No.1061〕"),
("Have [buy] a health insurance.","健康保険に入っている。〔No.1062〕"),
("Suffer financial hardship.","お金で苦労する（経済的苦難に苦しむ）。〔No.1063〕"),
("Live below the poverty line.","最低（貧困）水準以下の暮らしをする。〔No.1064〕"),
("The kitchen is in chaos.","台所がめちゃくちゃ（大混乱）だ。〔No.1065〕"),
("In isolation from society.","社会から孤立して。〔No.1066〕"),
("A chance for classroom interaction.","クラス内の交流のための機会。〔No.1067〕"),
("A wine-producing region.","ワインの生産地域。〔No.1068〕"),
("There is no proof that she was at home at that time.","彼女がその時家にいたという証拠はない。〔No.1069〕"),
("In principle.","原則的には。〔No.1070〕"),
("The origin of life on Earth.","地球上の生命の起源。〔No.1071〕"),
("A perfect setting for a picnic.","ピクニックにはうってつけの環境。〔No.1072〕"),
("Put up a monument.","記念碑を立てる。〔No.1073〕"),
("A mission to Saturn.","土星への（飛行）任務。〔No.1074〕"),
("The outline of the new project.","新企画の概要。〔No.1075〕"),
("A revolution in health care.","医療の大改革。〔No.1076〕"),
("Live on a pension.","年金で生活する。〔No.1077〕"),
("Our school was founded in 1918.","我が校は1918年に創立された。〔No.1078〕"),
("Sign a contract with the team.","そのチームと契約をする。〔No.1079〕"),
("Negotiate with kidnappers.","誘拐犯と交渉する。〔No.1080〕"),
("We all cooperated in preparing for the cultural festival.","私たち全員で協力して文化祭の準備をした。〔No.1081〕"),
("Restore peace in the Middle East.","中東での平和を回復する。〔No.1082〕"),
("The conservative party.","保守党。〔No.1083〕"),
("International affairs.","国際情勢。〔No.1084〕"),
("The Central Intelligence Agency (CIA).","アメリカ中央情報局。〔No.1085〕"),
("The city council.","市議会。〔No.1086〕"),
("The United Kingdom.","連合王国（＝英国）。〔No.1087〕"),
("The People's Republic of China.","中華人民共和国。〔No.1088〕"),
("The fall of the Roman Empire.","ローマ帝国の崩壊。〔No.1089〕"),
("Boycott the Security Council.","安全保障理事会をボイコットする。〔No.1090〕"),
("Give economic aid to developing countries.","発展途上国へ経済援助を行う。〔No.1091〕"),
("The government's educational reform.","政府の教育改革。〔No.1092〕"),
("The village mayor.","村長。〔No.1093〕"),
("The foreign minister.","外務大臣。〔No.1094〕"),
("Recent opinion polls.","最近の世論調査。〔No.1095〕"),
("Make a fortune in real estate.","不動産で一財産を作る。〔No.1096〕"),
("Intellectual property.","知的財産。〔No.1097〕"),
("Manage funds effectively.","資金を有効に活用する。〔No.1098〕"),
("Make money on the stock market.","株（式市場）でもうける。〔No.1099〕"),
("Go bankrupt owing one billion yen.","10億円の負債を抱えて倒産する。〔No.1100〕"),
("Manual labor.","肉体労働。〔No.1101〕"),
("Stop overworking.","働きすぎをやめる。〔No.1102〕"),
("A law firm.","法律事務所。〔No.1103〕"),
("A union member.","労働組合員。〔No.1104〕"),
("School administration.","学校運営。〔No.1105〕"),
("A secretary to the president.","社長秘書。〔No.1106〕"),
("The editor of the Japan News.","『ジャパンニュース』の編集長。〔No.1107〕"),
("The lawyer met a client yesterday.","その弁護士は依頼人と昨日会った。〔No.1108〕"),
("A colleague from work.","会社の同僚。〔No.1109〕"),
("On the basis of a new theory.","新しい理論を根拠にして。〔No.1110〕"),
("A key element of his success.","彼の成功の鍵となる要素。〔No.1111〕"),
("Fifty states constitute the USA.","50の州が米国を構成している。〔No.1112〕"),
("Unify the country.","国を統一する。〔No.1113〕"),
("Unite in fighting poverty.","貧困撲滅で団結する。〔No.1114〕"),
("Combine a diet with exercise.","ダイエットを運動と結びつける。〔No.1115〕"),
("a brutal murder","残忍な殺人。〔No.1116〕"),
("arrest a man for shoplifting","万引で男を逮捕する。〔No.1117〕"),
("Eric is in prison.","エリックは服役中（刑務所の中）だ。〔No.1118〕"),
("fight against vice on the streets","街の犯罪と戦う。〔No.1119〕"),
("the only witness to the murder","その殺人の唯一の目撃者。〔No.1120〕"),
("a conflict between two parties","2つの政党の間の対立。〔No.1121〕"),
("The bullet passed through his right leg.","弾丸は彼の右足を貫通した。〔No.1122〕"),
("The wound still hurts.","傷がまだ痛む。〔No.1123〕"),
("her triumph in the election","選挙での彼女の勝利。〔No.1124〕"),
("use military force","軍事力を使う。〔No.1125〕"),
("a strategy for winning the game","その試合に勝つための戦略。〔No.1126〕"),
("The Normans conquered England in 1066.","1066年, ノルマン人はイングランドを征服した。〔No.1127〕"),
("learn karate to defend myself","自分の身を守るために空手を習う。〔No.1128〕"),
("quarrel with a neighbor","近所の人と口論する。〔No.1129〕"),
("The pirates robbed the ship.","海賊はその船を襲った。〔No.1130〕"),
("Diana was deprived of her civil rights.","ダイアナは公民権を剥奪された。〔No.1131〕"),
("The city was devastated by a big earthquake.","その都市は大地震で壊滅した（壊滅させられた）。〔No.1132〕"),
("Steep stairs can be a hazard for elderly people.","急な階段は高齢者にとって危険（なもの）になり得る。〔No.1133〕"),
("a war orphan","戦争孤児。〔No.1134〕"),
("a breast-fed infant","母乳で育てられた幼児。〔No.1135〕"),
("free the slaves","奴隷を解放する。〔No.1136〕"),
("an old acquaintance","昔からの知り合い。〔No.1137〕"),
("a wine merchant","ワイン店を営む人。〔No.1138〕"),
("Parking spaces are for residents only.","駐車スペースは居住者専用です。〔No.1139〕"),
("passengers and crew on the plane.","飛行機の乗客と乗組員。〔No.1140〕"),
("commute to work by train","電車を使って通勤する。〔No.1141〕"),
("Police chased the stolen car.","警察はその盗難車を追いかけた。〔No.1142〕"),
("overtake the truck ahead of us","前のトラックを追い越す。〔No.1143〕"),
("Childhood memories fade as time passes.","時が過ぎるとともに子どものころの記憶は薄れる。〔No.1144〕"),
("Salt dissolves in hot water.","塩が熱湯に溶ける。〔No.1145〕"),
("float on the sea","海面に浮かぶ。〔No.1146〕"),
("Oil does not sink in water.","油は水に沈まない。〔No.1147〕"),
("Stock prices are fluctuating wildly.","株価は乱高下している（激しく変動している）。〔No.1148〕"),
("Blood circulates through the body.","血液は体内を循環する。〔No.1149〕"),
("transform the old house into a restaurant","古い家を改造してレストランにする。〔No.1150〕"),
("a paradigm shift","パラダイムシフト（理論的枠組の転換）。〔No.1151〕"),
("genetically modified (GM) foods","遺伝子組み換え（遺伝子上修正された）食品。〔No.1152〕"),
("revise the educational system","教育制度を改正する。〔No.1153〕"),
("emergency vehicles","緊急車両。〔No.1154〕"),
("take a voyage around the world","世界一周の船旅をする。〔No.1155〕"),
("confirm the booking","予約を確認する。〔No.1156〕"),
("Please ensure that you leave nothing behind.","忘れ物をしないように気をつけ（確実にし）てください。〔No.1157〕"),
("address an environmental problem","環境問題に取り組む。〔No.1158〕"),
("a new approach to teaching languages","言語教育への新たな取り組み方。〔No.1159〕"),
("resolve the matter on my own","自分でその問題を解決する。〔No.1160〕"),
("settle the dispute","その紛争を解決する。〔No.1161〕"),
("warn him to follow the rules","規則に従うよう彼に警告する。〔No.1162〕"),
("The older kids forced him to shoplift.","年上の子達が彼に万引するよう強いた。〔No.1163〕"),
("boost the team's morale","チームの士気を高める。〔No.1164〕"),
("appeal to our emotions","私たちの感情に訴えかける。〔No.1165〕"),
("A fish leaped out of a pond.","魚が池の中から跳び出た。〔No.1166〕"),
("grab him by the neck","彼の首根っこをつかむ。〔No.1167〕"),
("seize her by the shoulder","彼女の肩をつかむ。〔No.1168〕"),
("cast a net around a school of tuna.","マグロの群れに網を投げ入れる。〔No.1169〕"),
("slap him in the face.","彼の顔を平手打ちする。〔No.1170〕"),
("stun him with a blow.","殴って彼を気絶させる。〔No.1171〕"),
("a beautifully illuminated castle.","美しく照らされた城。〔No.1172〕"),
("mend a tear in my jacket.","上着のほころびを修繕する。〔No.1173〕"),
("react angrily to the news.","その知らせに怒りの反応を示す（怒って反応する）。〔No.1174〕"),
("endure years of suffering.","長年の苦労に耐える。〔No.1175〕"),
("encounter a bear in the woods.","森の中でクマに遭遇する。〔No.1176〕"),
("neglect my duty.","義務を怠る。〔No.1177〕"),
("undergo cosmetic surgery.","美容整形手術を受ける。〔No.1178〕"),
("trace a missing person.","行方不明の人を捜し出す。〔No.1179〕"),
("conduct a survey.","調査を行う。〔No.1180〕"),
("emerge from the darkness.","暗闇から現れる。〔No.1181〕"),
("How will that drama unfold?","そのドラマはどのように展開するのだろう。〔No.1182〕"),
("This word derives from Latin.","この単語はラテン語に由来する。〔No.1183〕"),
("Every child possesses a range of abilities.","どの子も多彩な才能を所有している。〔No.1184〕"),
("retain her beauty.","彼女の美しさを保持する。〔No.1185〕"),
("secure a window seat.","窓側の席を確保する。〔No.1186〕"),
("The monkey that escaped from the zoo was captured.","動物園から逃げたサルが確保された。〔No.1187〕"),
("an inquiry into his background.","彼の身元調査。〔No.1188〕"),
("specialize in business administration.","経営学を専攻する。〔No.1189〕"),
("the first semester.","前期。〔No.1190〕"),
("Biology is the scientific study of living things.","生物学は生物の科学的な研究です。〔No.1191〕"),
("the ecology of jellyfish.","クラゲの生態。〔No.1192〕"),
("the philosophy of Aristotle.","アリストテレスの哲学。〔No.1193〕"),
("a great discovery in geography.","地理上の偉大な発見。〔No.1194〕"),
("a book of psychology.","心理学の本。〔No.1195〕"),
("educational institutions.","教育機関。〔No.1196〕"),
("live in a dormitory.","寮生活をする。〔No.1197〕"),
("the notion that smoking is cool.","タバコを吸うのは格好いいという考え。〔No.1198〕"),
("the concept of time.","時間の概念。〔No.1199〕"),
("sit in Zen meditation.","座禅（座って行う瞑想）を組む。〔No.1200〕"),
("an insight into human nature.","人間性への洞察。〔No.1201〕"),
("broaden my outlook on life.","人生観を広げる。〔No.1202〕"),
("act with caution.","用心して行動する。〔No.1203〕"),
("against my will.","意思に反して。〔No.1204〕"),
("consider other options.","ほかの選択肢を考える。〔No.1205〕"),
("have a terrible nightmare.","ひどい悪夢を見る。〔No.1206〕"),
("have a keen eye for talent.","才能を見る(鋭い)目がある。〔No.1207〕"),
("cherish the memories of that day.","その日の思い出を大切にする。〔No.1208〕"),
("intend to see the movie.","その映画を見るつもりだ。〔No.1209〕"),
("The police suspect that she stole the money.","警察は彼女がそのお金を盗んだのではないかと思っている。〔No.1210〕"),
("I could not comprehend what had happened.","私は何が起きたのかが理解できなかった。〔No.1211〕"),
("I am afraid you misunderstand me.","君は私のことを誤解しているようだね。〔No.1212〕"),
("identify the fingerprints.","その指紋を特定する。〔No.1213〕"),
("Parking is not permitted here.","ここでは駐車は許可されていません。〔No.1214〕"),
("acknowledge the need for change.","変化の必要を認める。〔No.1215〕"),
("trim the fat off the meat.","肉から脂身を切り落とす。〔No.1216〕"),
("His name was omitted from the list.","彼の名前は一覧には入っていなかった。〔No.1217〕"),
("cultivate the land.","土地を耕す。〔No.1218〕"),
("sit in the shade.","日陰に座る。〔No.1219〕"),
("feel a sea breeze.","海のそよ風を感じる。〔No.1220〕"),
("the Sahara desert.","サハラ砂漠。〔No.1221〕"),
("the dense fog cleared.","濃い霧が晴れた。〔No.1222〕"),
("Moisture is essential for keeping your skin youthful.","湿り気は皮膚を若々しく保つのに不可欠だ。〔No.1223〕"),
("My shirt is still damp.","シャツがまだ湿っている。〔No.1224〕"),
("a path to the mountaintop.","山頂への道。〔No.1225〕"),
("a hiking trail.","ハイキングコース。〔No.1226〕"),
("a mountain stream.","谷川。〔No.1227〕"),
("the tide is in.","満潮である。〔No.1228〕"),
("the natural habitat of sea lions.","アシカの自然生息地。〔No.1229〕"),
("the ozone layer.","オゾン層。〔No.1230〕"),
("that volcano erupted last year.","その火山は昨年噴火した。〔No.1231〕"),
("work in a coal mine.","炭鉱で働く。〔No.1232〕"),
("the sweet scent of perfume.","香水の甘い香り。〔No.1233〕"),
("Wheat is the chief crop in this area.","小麦はこの地域の主な作物です。〔No.1234〕"),
("One swallow does not make a summer.","早合点は禁物。〔No.1235〕"),
("an insect caught in a spider's web.","クモの巣にかかった虫。〔No.1236〕"),
("cattle, sheep, and other domestic animals.","ウシ, ヒツジ, とそれ以外の家畜。〔No.1237〕"),
("a lost kitten.","迷子の子ネコ。〔No.1238〕"),
("microscopic organisms.","微生物。〔No.1239〕"),
("the wildlife of the Galapagos Islands.","ガラパゴス諸島の野生生物。〔No.1240〕"),
("the lion, the king of beasts.","百獣の王ライオン。〔No.1241〕"),
("great apes.","大型の類人猿。〔No.1242〕"),
("Whales are classified as mammals.","クジラは哺乳類に分類されている。〔No.1243〕"),
("a flock of crows.","カラスの一群。〔No.1244〕"),
("Rabbits breed all year round.","ウサギは年中繁殖する。〔No.1245〕"),
("The virus can reproduce rapidly.","ウイルスは急速に繁殖する。〔No.1246〕"),
("The eggs hatched in thirty days.","その卵は30日でかえった。〔No.1247〕"),
("a diligent student.","勤勉な学生。〔No.1248〕"),
("his noble character.","彼の高潔な人格。〔No.1249〕"),
("Society's elite tend to be arrogant.","エリートは傲慢になりがちだ。〔No.1250〕"),
("Rabbits are timid animals.","ウサギは臆病な動物です。〔No.1251〕"),
("a humble character.","控えめな性格の人。〔No.1252〕"),
("People are inclined to judge someone by their appearance.","人は人を外見で判断する傾向にある。〔No.1253〕"),
("the country's principal export.","その国の主な輸出品。〔No.1254〕"),
("the prime minister of Japan.","日本の首相。〔No.1255〕"),
("Cars are indispensable to a rural lifestyle.","車は田舎の生活に不可欠だ。〔No.1256〕"),
("The ship was in grave danger.","その船は重大な危機に瀕していた。〔No.1257〕"),
("a definite answer.","明確な答え。〔No.1258〕"),
("It was evident that he was not telling the truth.","彼が真実を言っていないのは明らかだった。〔No.1259〕"),
("The statue stood in a prominent position.","その像は目立った場所にあった。〔No.1260〕"),
("a marked change in his behavior.","彼の行動の際立った変化。〔No.1261〕"),
("the grace of the dancer.","そのダンサーの優美さ。〔No.1262〕"),
("have enormous charm.","大きな魅力がある。〔No.1263〕"),
("make a rational decision.","理性的な決定をする。〔No.1264〕"),
("a magnificent view of Mt. Fuji.","富士山の壮大な眺め。〔No.1265〕"),
("Your bike is superior to mine.","君の自転車は私のより優れている。〔No.1266〕"),
("a loyal supporter of the team.","そのチームの忠実な支持者。〔No.1267〕"),
("That fish looks horrible but tastes terrific.","その魚は見た目はひどいが, 味はすばらしい。〔No.1268〕"),
("lead a miserable life.","悲惨な生活を送る。〔No.1269〕"),
("a dull TV program.","退屈なテレビ番組。〔No.1270〕"),
("leave a nasty taste in my mouth.","後味が悪い。〔No.1271〕"),
("Have you gone insane?","正気を失ったか？〔No.1272〕"),
("I believe that lowering the voting age is a ridiculous idea.","選挙年齢を引き下げるのは馬鹿げた考えだと思う。〔No.1273〕"),
("That clothing shop is notorious for its poor quality.","その服屋は低品質で悪名高い。〔No.1274〕"),
("an evil spirit.","悪霊。〔No.1275〕"),
("make a mess of the room.","部屋をめちゃくちゃにする。〔No.1276〕"),
("a vast palace.","広大な宮殿。〔No.1277〕"),
("a huge elephant.","巨大なゾウ。〔No.1278〕"),
("a tiny hole in the wall.","壁に開いたとても小さな穴。〔No.1279〕"),
("a perfect sphere.","完全な球体。〔No.1280〕"),
("a strip of paper.","細長い紙片。〔No.1281〕"),
("internal organs.","内蔵。〔No.1282〕"),
("a mobile society.","流動的な社会。〔No.1283〕"),
("climb steep stone steps.","急な石段を上る。〔No.1284〕"),
("intense pressure to succeed.","成功しなければという強烈な重圧。〔No.1285〕"),
("an abstract work of art.","抽象的な芸術作品。〔No.1286〕"),
("an inevitable result.","避けられない結果。〔No.1287〕"),
("a demilitarized neutral zone.","非武装中立地帯。〔No.1288〕"),
("have a brief conversation with her.","彼女と少し話をする。〔No.1289〕"),
("a potential customer.","将来の顧客。〔No.1290〕"),
("a tough soldier.","たくましい兵士。〔No.1291〕"),
("secondhand bookstores.","古本屋。〔No.1292〕"),
("Peter resembles his father in every way.","ピーターはあらゆる点で父親に似ている。〔No.1293〕"),
("Roger is not stupid; on the contrary, he is a genius.","ロジャーはバカではない. それどころか天才だ。〔No.1294〕"),
("fall into the same category.","同じ範疇に属する。〔No.1295〕"),
("limit the scope of our activities.","私たちの活動範囲を制限する。〔No.1296〕"),
("Go back to your respective rooms.","それぞれの部屋に戻りなさい。〔No.1297〕"),
("a given period of time.","ある一定の期間。〔No.1298〕"),
("a custom peculiar to Japan.","日本固有の習慣。〔No.1299〕"),
("plenty of space for kids to play.","子どもが遊ぶのに十分な広さ。〔No.1300〕"),
("sufficient food for everyone.","全員に行き渡る十分な食料。〔No.1301〕"),
("on numerous occasions.","多くの場面で。〔No.1302〕"),
("a substantial breakfast.","ボリュームのある朝食。〔No.1303〕"),
("The family is the smallest unit of society.","家族は社会の最小の単位だ。〔No.1304〕"),
("two meters in diameter.","直径２メートル。〔No.1305〕"),
("Quality is more important than quantity.","量より質が重要だ。〔No.1306〕"),
("a vitamin deficiency.","ビタミンの欠乏。〔No.1307〕"),
("lessen the risk of human error.","人為的ミスのリスクを軽減する。〔No.1308〕"),
("the dawn of a new era.","新たな時代の幕明け。〔No.1309〕"),
("have a Q&A session.","質疑応答の時間をとる。〔No.1310〕"),
("begin to speak after a long pause.","長い休止の後に話し始める。〔No.1311〕"),
("My wife was shopping. Meanwhile, I was waiting in the car.","妻は買い物をしていた. その間, 私は車で待っていた。〔No.1312〕"),
("a souvenir of my trip to France.","(自分のための)フランス旅行の記念品。〔No.1313〕"),
("a brick house.","れんが造りの家。〔No.1314〕"),
("a ten-foot bamboo pole.","10フィートの竹ざお。〔No.1315〕"),
("sew with silk thread.","絹糸で縫う。〔No.1316〕"),
("a child who loves mischief.","いたずらが好きな子ども。〔No.1317〕"),
("the superstition that rooms numbered 4 are unlucky.","４号室は不吉であるという迷信。〔No.1318〕"),
("I need more storage space.","もっと多くの貯蔵スペースが必要だ。〔No.1319〕"),
("take shelter from the rain.","雨宿りする。〔No.1320〕"),
("Delicate garments should be handwashed.","傷みやすい衣類は手洗いしなければならない。〔No.1321〕"),
("put the forms in the filing cabinet.","書類整理棚にその書類をしまう。〔No.1322〕"),
("learn the recipe for this stew.","このシチューの調理法を学ぶ。〔No.1323〕"),
("look over a travel brochure.","旅行のパンフレットに目を通す。〔No.1324〕"),
("Free shipping anywhere within Japan.","日本国内ならどこでも送料無料。〔No.1325〕"),
("fetch the children from the kindergarten.","幼稚園に子どものお迎えに行く。〔No.1326〕"),
("distribute handouts.","プリントを配る。〔No.1327〕"),
("Let's meet at the Hachiko statue at ten.","10時にハチ公の所で会いましょう。〔No.1328〕"),
("The Thinker is a sculpture by Rodin.","「考える人」はロダンの彫刻だ。〔No.1329〕"),
("carve her initials into a tree.","木に彼女のイニシャルを刻む。〔No.1330〕"),
("a myth about the safety of nuclear power plants.","原子力発電所の安全神話。〔No.1331〕"),
("traditional crafts of this region.","この知己の伝統的な民芸品。〔No.1332〕"),
("an academic journal.","学会誌。〔No.1333〕"),
("go to heaven.","天国に行く。〔No.1334〕"),
("attend a funeral.","葬式に出席する。〔No.1335〕"),
("the souls of the dead.","死者の魂。〔No.1336〕"),
('The priest crossed himself, saying "Amen."',"神父は「アーメン」と言いながら十字を切った。〔No.1337〕"),
("the Buddhist faith.","仏教(仏教の信仰)。〔No.1338〕"),
("Cows are sacred animals to the Hindus.","ウシはヒンドゥー教徒にとって神聖な動物だ。〔No.1339〕"),
("a divine message.","神のお告げ。〔No.1340〕"),
('The meaning of "charge" depends on the context.',"chargeの意味は文脈によって決まる。〔No.1341〕"),
("a singular form.","単数形。〔No.1342〕"),
("have a good command of Chinese.","中国語が達者だ。〔No.1343〕"),
("Take Lucy, for instance.","たとえばルーシーを例に取り上げてみよう。〔No.1344〕"),
("a first-person narrative.","一人称の語り。〔No.1345〕"),
("a book of English dialogues.","英会話集。〔No.1346〕"),
("grammar and usage.","文法と語法。〔No.1347〕"),
("He speaks English with a strong German accent.","彼は強いドイツ語訛の英語を話す。〔No.1348〕"),
("have a violent temper.","荒い気性だ。〔No.1349〕"),
("feel great affection for my grandfather.","祖父に強い愛情を感じる。〔No.1350〕"),
("with great enthusiasm.","やる気満々で。〔No.1351〕"),
("have a passion for music.","音楽への情熱を持っている。〔No.1352〕"),
("Allan was red with rage.","アランは怒りで真っ赤になっていた。〔No.1353〕"),
("the family's joys and sorrows.","家族の喜びと悲しみ。〔No.1354〕"),
("incredible news.","信じられない知らせ。〔No.1355〕"),
("chuckle at the memory.","思い出し笑いをする。〔No.1356〕"),
("Stop teasing him!","彼をからかうのをやめて。〔No.1357〕"),
("No one dared to argue.","言い争う勇気がある人はいなかった。〔No.1358〕"),
("He was willing to help me.","彼は嫌がらずに私を助けてくれた。〔No.1359〕"),
("scream in horror.","恐怖で悲鳴を上げる。〔No.1360〕"),
("soothe a crying baby.","泣いている赤ん坊をなだめる。〔No.1361〕"),
("The sunny weather tempted me to go out.","晴れていたので, 外出したくなった。〔No.1362〕"),
("scold my son for making a mess of his room.","部屋を散らかしたことに対して息子を叱る。〔No.1363〕"),
("The dictator oppressed his people for many years.","その独裁者は長年にわたり人民を虐げた。〔No.1364〕"),
("threaten to report it to the police.","それを警察に通報すると脅す。〔No.1365〕"),
("I was delighted to hear the news.","その知らせを聞いて喜んだ。〔No.1366〕"),
("We were astonished at the unexpected news.","私たちはその予期せぬ知らせに仰天した。〔No.1367〕"),
("mourn the death of the Pope.","ローマ教皇の死を悼む。〔No.1368〕"),
("His jokes offended Irish people.","彼の冗談はアイルランドの人たちを怒らせた。〔No.1369〕"),
("The rainy weather depressed me.","雨天ですっかり気がめいってしまった。〔No.1370〕"),
("sacrifice my family for my career.","仕事のために家族を犠牲にする。〔No.1371〕"),
("buy a ring on impulse.","指輪を衝動買いする。〔No.1372〕"),
("Gambling is a curse on his life.","ギャンブルは彼にとって災いの元だ。〔No.1373〕"),
("a key factor in his success.","彼の成功のかぎとなる要因。〔No.1374〕"),
("Diabetes often stems from an unhealthy lifestyle.","糖尿病は不健康な生活習慣が原因でなることが多い。〔No.1375〕"),
("trigger an allergic reaction.","アレルギー反応を引き起こす。〔No.1376〕"),
("have serious consequences.","深刻な結果を招く。〔No.1377〕"),
("The production process is fully automated.","生産過程は完全にオートメーション化されている。〔No.1378〕"),
("resort to violence.","暴力に訴える。〔No.1379〕"),
("a new formula.","新しい方法。〔No.1380〕"),
("the procedure for applying for a passport.","パスポート申請の手順。〔No.1381〕"),
("have an extraordinary memory.","並外れた記憶力を有する。〔No.1382〕"),
("in absolute silence.","静粛に(絶対的な沈黙で)。〔No.1383〕"),
("an overwhelming majority.","圧倒的多数。〔No.1384〕"),
("This old custom has vanished altogether.","この古い習慣は完全になくなった。〔No.1385〕"),
("The entire audience stood up to cheer.","観客は総立ちで歓声を上げた。〔No.1386〕"),
("a tremendous reaction from the crowd.","群衆からのすさまじい反応。〔No.1387〕"),
("the sheer number of visitors.","非常に多くの訪問者。〔No.1388〕"),
("a partial success.","部分的な成功。〔No.1389〕"),
("barely enough food to go around.","かろうじて行き渡るだけの食物。〔No.1390〕"),
("have a slight fever.","微熱がある。〔No.1391〕"),
("approximately one million yen.","およそ100万円。〔No.1392〕"),
("Buses run frequently to the airport.","空港へのバスの便は頻繁にあります。〔No.1393〕"),
("Apparently, Jim's a good cook.","ジムはどうやら料理が上手いようだ。〔No.1394〕"),
("Jill is merely a child.","ジルは子どもにすぎない。〔No.1395〕"),
("It was snowing; nevertheless, we continued to work in the fields.","雪が降っていた. それにもかかわらず, 私たちは畑仕事をし続けた。〔No.1396〕"),
("escape from the box somehow.","何らかの方法でその箱から脱出する。〔No.1397〕"),
("This problem is very simple, and thus easy to solve.","この問題はとても単純なので, それゆえ, 解くのは簡単だ。〔No.1398〕"),
("I know him pretty well.","彼のことはまあまあ知ってるよ。〔No.1399〕"),
("The soccer game was held, despite the heavy rain.","大雨にもかかわらず, そのサッカーの試合は行われた。〔No.1400〕"),
("assert that French cuisine is the world's best.","フランス料理は世界で一番だと主張する。〔No.1401〕"),
("defy the teacher.","教師に反抗する。〔No.1402〕"),
("condemn racial discrimination.","人種差別を非難する。〔No.1403〕"),
("The two statements contradict each other.","その2つの供述は互いに矛盾している。〔No.1404〕"),
("cite the example of Sweden.","スウェーデンの例を引き合いに出す。〔No.1405〕"),
("use several examples to illustrate my point.","私の主張を説明するためにいくつかの例を用いる。〔No.1406〕"),
("articulate my ideas to the boss.","上司に自分の考えをはっきり述べる。〔No.1407〕"),
("an advocate of disarmament.","軍縮提唱者。〔No.1408〕"),
("a breakthrough in the treatment of cancer.","がん治療の飛躍的進歩。〔No.1409〕"),
("Darwin's theory of evolution.","ダーウィンの進化論。〔No.1410〕"),
("the speed of technological innovation.","技術革新のスピード。〔No.1411〕"),
("a space probe on its way to Mars.","火星に向かう宇宙探査機。〔No.1412〕"),
("drink plenty of fluids.","十分な水分を補給する。〔No.1413〕"),
("an elementary particle.","素粒子(初歩的な粒子)。〔No.1414〕"),
("in zero gravity.","無重力で。〔No.1415〕"),
("trade friction between the two countries.","両国間の貿易摩擦。〔No.1416〕"),
("a gene mutation.","遺伝子の突然変異。〔No.1417〕"),
("Mendel's laws of heredity.","メンデルの遺伝の法則。〔No.1418〕"),
("test a hypothesis.","仮説を検証する。〔No.1419〕"),
("butterfly specimens.","チョウの標本。〔No.1420〕"),
("look at microbes under a microscope.","顕微鏡で微生物を見る。〔No.1421〕"),
("new discoveries in the realm of biochemistry.","生化学の領域における新たな発見。〔No.1422〕"),
("acidic rain.","酸性雨。〔No.1423〕"),
("toxic substances.","有害物質。〔No.1424〕"),
("The skull and crossbones is a warning symbol for poison.","ドクロと交差した骨は毒であることを警告する象徴だ。〔No.1425〕"),
("strive to cut costs.","経費削減に努力する。〔No.1426〕"),
("despite my best endeavors.","最大の努力にもかかわらず。〔No.1427〕"),
("manufacture computer parts.","コンピュータ部品を製造する。〔No.1428〕"),
("yield a large profit.","大きな利益を生む。〔No.1429〕"),
("duplicate a document.","書類を複製する。〔No.1430〕"),
("assemble thousands of cars.","何千台もの車を組み立てる。〔No.1431〕"),
("mold a chocolate bar.","板チョコを作る。〔No.1432〕"),
('The term "black hole" was coined in 1969.',"「ブラックホール」という言葉は1969年に作られた。〔No.1433〕"),
("emergency physicians.","救急医。〔No.1434〕"),
("A veterinarian is a doctor for animals.","獣医とは動物の医師のことです。〔No.1435〕"),
("gene therapy.","遺伝子治療。〔No.1436〕"),
("get an injection.","注射をしてもらう。〔No.1437〕"),
("Take a dose of three pills after every meal.","毎食後3錠(の服用量)ずつ服用してください。〔No.1438〕"),
("improve public sanitation.","公衆衛生を改善する。〔No.1439〕"),
("germs and dust.","細菌とほこり。〔No.1440〕"),
("remove a brain tumor.","脳腫瘍を除去する。〔No.1441〕"),
("Having my tooth drilled was agony.","歯を削られるのは苦痛だった。〔No.1442〕"),
("eating disorders.","摂食障害。〔No.1443〕"),
("die from the plague.","疫病で死ぬ。〔No.1444〕"),
("a flu epidemic.","インフルエンザの大流行。〔No.1445〕"),
("receive a medical diagnosis.","医師の診断を受ける。〔No.1446〕"),
("the infant mortality rate.","幼児死亡率。〔No.1447〕"),
("chronic back pain.","慢性の腰痛。〔No.1448〕"),
("Most adults are immune to measles.","たいていの大人ははしかに免疫がある。〔No.1449〕"),
("choke on a rice cake.","もにをのどに詰まらせる。〔No.1450〕"),
("Snow paralyzed public transportation.","雪のため交通機関は麻痺した。〔No.1451〕"),
("My dog was infected with a virus.","うちのイヌはウイルスに感染した。〔No.1452〕"),
("prescribe some sleeping pills.","睡眠薬を処方する。〔No.1453〕"),
("something easy to digest.","消化によい(食べ)物。〔No.1454〕"),
("weary hostages.","疲労した人質たち。〔No.1455〕"),
("nourish the skin.","皮膚に栄養を与える。〔No.1456〕"),
("the flesh of animals.","動物の肉。〔No.1457〕"),
("have a kidney transplant.","腎臓移植手術を受ける。〔No.1458〕"),
("You have good posture.","君は姿勢がいい。〔No.1459〕"),
("have long limbs.","手足が長い。〔No.1460〕"),
("walk erect.","直立歩行する。〔No.1461〕"),
("The airline compensated us for our lost baggage.","航空会社は紛失した私たちの手荷物に対して(私達に)補償をした。〔No.1462〕"),
("Lisbon flourished as a trading center.","リスボンは貿易の中心地として栄えた。〔No.1463〕"),
("child abuse.","児童の虐待。〔No.1464〕"),
("a bias against single mothers.","シングルマザーに対する偏見。〔No.1465〕"),
("gain wealth and prestige.","富と名声を得る。〔No.1466〕"),
("a privilege of the rich.","金持ちの特権。〔No.1467〕"),
("accomplish the feat of climbing Mt. Everest twice.","エベレスト山に2回登るという偉業を達成する。〔No.1468〕"),
("choose a death with dignity.","尊厳死を選ぶ。〔No.1469〕"),
("lead a life of virtue.","徳のある生活を送る。〔No.1470〕"),
("a pastime of the aristocracy.","貴族の娯楽。〔No.1471〕"),
("a rebel army.","反乱軍。〔No.1472〕"),
("a clever marketing scheme.","巧妙なマーケティング計画。〔No.1473〕"),
("a convention for anime fans.","アニメファンの大会。〔No.1474〕"),
("a five-day conference on global warming.","地球温暖化に関する5日間の会議。〔No.1475〕"),
("get married without our parents' consent.","親の同意なしで結婚する。〔No.1476〕"),
("a diplomatic initiative for the Middle East.","中東外交に対する新構想。〔No.1477〕"),
("the province of Quebec.","(カナダの)ケベック州。〔No.1478〕"),
("a former British colony.","旧英国植民地。〔No.1479〕"),
("a tribe living on the frontier.","辺境に暮らす部族。〔No.1480〕"),
("a tribe native to Hokkaido.","北海道の先住民。〔No.1481〕"),
("a primitive society.","原始社会。〔No.1482〕"),
("ethnic minorities.","少数民族。〔No.1483〕"),
("get an anonymous phone call.","非通知の電話を受ける。〔No.1484〕"),
("media coverage on the case.","その事件に関するマスコミ報道。〔No.1485〕"),
("fulfill social obligations.","社会的責務を果たす。〔No.1486〕"),
("Arranged marriages were the norm in the past.","お見合い結婚は昔はふつうのことだった。〔No.1487〕"),
("file a patent.","特許を申請する。〔No.1488〕"),
("She is entitled to receive a pension.","彼女は年金をもらう権利がある。〔No.1489〕"),
("donate ten million yen to my old school.","母校に1,000万円を寄付する。〔No.1490〕"),
("pressure to conform to the group.","集団に合わせる重圧。〔No.1491〕"),
("comply with the regulations.","その規則を遵守する。〔No.1492〕"),
("new legislation on the sale of guns.","銃販売の新たな法律。〔No.1493〕"),
("the witness's testimony.","目撃者の証言。〔No.1494〕"),
("a trial by jury.","陪審裁判。〔No.1495〕"),
("a legitimate reason for being late.","遅刻に対する正当な理由。〔No.1496〕"),
("enforce the law.","その法律を施行する。〔No.1497〕"),
("the Nazi regime in Germany.","独のナチ政権。〔No.1498〕"),
("the government bureaucracy.","政府の官僚制度。〔No.1499〕"),
("corruption of politicians.","政治家の腐敗。〔No.1500〕"),
("the tyranny of the government.","政府の圧政。〔No.1501〕"),
("impose sanctions on that country.","その国に制裁を加える。〔No.1502〕"),
("take a census every five years.","5年ごとに国勢調査を行う。〔No.1503〕"),
("a candidate for the Nobel Prize.","ノーベル賞の候補者。〔No.1504〕"),
("Senator Kennedy.","ケネディ上院議員。〔No.1505〕"),
("the Japanese ambassador to the USA.","駐米日本大使。〔No.1506〕"),
("an experienced diplomat.","経験豊かな外交官。〔No.1507〕"),
("refugee camps in Syria.","シリアの難民キャンプ。〔No.1508〕"),
("the issue of the Northern territories.","北方領土問題。〔No.1509〕"),
("an important ally.","重要な同盟国。〔No.1510〕"),
("the Federal Bureau of Investigation (FBI).","(米国)連邦捜査局。〔No.1511〕"),
("the feudal system in Europe.","ヨーロッパの封建制度。〔No.1512〕"),
("the Chinese Communist Party.","中国共産党。〔No.1513〕"),
("real estate.","不動産。〔No.1514〕"),
("She is a great asset to the company.","彼女は会社にとって大きな財産だ。〔No.1515〕"),
("an increase in tax revenues.","税収の増加。〔No.1516〕"),
("the USA trade deficit with Japan.","アメリカの対日貿易赤字。〔No.1517〕"),
("commerce and industry.","商工業。〔No.1518〕"),
("Australian steel output.","オーストラリアの鋼鉄の生産力。〔No.1519〕"),
("receive subsidies from the government.","政府から補助金を受け取る。〔No.1520〕"),
("a monopoly on oil sales.","石油販売の独占権。〔No.1521〕"),
("online banking transactions.","オンライン銀行取り引き。〔No.1522〕"),
("a corporation with 2000 employees.","従業員2000人を抱える企業。〔No.1523〕"),
("a large enterprise.","大企業。〔No.1524〕"),
("embark on a joint venture.","合弁事業に乗り出す。〔No.1525〕"),
("the United Nations headquarters.","国連本部。〔No.1526〕"),
("medical personnel.","医療関係職員。〔No.1527〕"),
("a wide range of expertise.","幅広い専門的知識。〔No.1528〕"),
("meet our monthly quota.","毎月のノルマを達成する。〔No.1529〕"),
("store goods in a warehouse.","商品を倉庫に保管する。〔No.1530〕"),
("toil in the fields.","畑で骨折って働く。〔No.1531〕"),
("undertake a big project.","大きな事業を引き受ける。〔No.1532〕"),
("The two departments have merged.","その２つの部署が合併した。〔No.1533〕"),
("The USA comprises 50 states.","アメリカは50の州から成っている。〔No.1534〕"),
("aircraft engine components.","飛行機のエンジン部品。〔No.1535〕"),
("the basic framework of his theory.","彼の理論の基本的枠組。〔No.1536〕"),
("a prototype of the device.","その装置の原型。〔No.1537〕"),
("the political dimensions of the meeting.","その会議の政治的な側面。〔No.1538〕"),
("write in the margins.","余白に書き込む。〔No.1539〕"),
("Your arguments are not relevant to this discussion.","君の主張はこの議論と関連がない。〔No.1540〕"),
("His blood type is not compatible with hers.","彼の血液型は彼女のと適合しない。〔No.1541〕"),
("The letters of the alphabet do not correspond exactly to their sounds.","アルファベットの文字は厳密には音に一致していない。〔No.1542〕"),
("identical twins.","一卵性双生児（うり二つの双子）。〔No.1543〕"),
("a bonus equivalent to 3 months' salary.","月給3か月分の（3か月分に等しい）ボーナス。〔No.1544〕"),
("a random choice.","無作為な選択。〔No.1545〕"),
("A warrior fights for glory, while a soldier fights for justice.","戦士は栄光のために戦い、一方、兵士は正義のために戦う。〔No.1546〕"),
("assaults on station staff.","駅員たちに対する暴行。〔No.1547〕"),
("send troops to the disaster-affected area.","被災地に軍を派遣する。〔No.1548〕"),
("remain silent even under torture.","拷問を受けても口を割らない。〔No.1549〕"),
("settle the border disputes.","国境問題（紛争）を解決する。〔No.1550〕"),
("crack down on a riot.","暴動を鎮圧する。〔No.1551〕"),
("Many Jews went to the USA to escape persecution in Europe.","ヨーロッパでの迫害を逃れるため、多くのユダヤ人が米国へ行った。〔No.1552〕"),
("Donald is dominated by his wife.","ドナルドは奥さんの尻に敷かれている（奥さんに支配されている）。〔No.1553〕"),
("Thousands of seals are slaughtered there every year.","そこでは毎年何千頭ものアザラシが殺されている。〔No.1554〕"),
("suppress a rebellion.","反乱を鎮圧する。〔No.1555〕"),
("surrender to the enemy.","敵に降伏する。〔No.1556〕"),
("have to contend with difficulties.","困難と戦わなくてはならない。〔No.1557〕"),
("She is hostile to me.","彼女は私に対して敵対的だ。〔No.1558〕"),
("the wreck of the Titanic.","タイタニック号の残骸。〔No.1559〕"),
("the death toll from the earthquake.","その地震の死亡者数。〔No.1560〕"),
("disrupt the train schedule.","鉄道のダイヤを乱す。〔No.1561〕"),
("The music interfered with my studies.","その音楽は私の勉強の邪魔になった。〔No.1562〕"),
("Some difficulties confronted us.","私たちは困難に直面していた（困難が私たちに立ちふさがった）。〔No.1563〕"),
("hinder economic growth.","経済成長を妨げる。〔No.1564〕"),
("A bomb exploded in the warehouse.","倉庫で爆弾が爆発した。〔No.1565〕"),
("A bus collided (head-on) with a truck.","バスがトラックと（正面）衝突した。〔No.1566〕"),
("I gave my nephew 5000 yen as a New Year's present.","甥にお年玉として5000円をあげた。〔No.1567〕"),
("the employees and their spouses.","従業員とその配偶者。〔No.1568〕"),
("the heir to the throne.","王位の継承者。〔No.1569〕"),
("a folk tale.","民話（民族固有の説話）。〔No.1570〕"),
("It is easy to relax among your peers.","同等の者同士で集まると気が楽だ。〔No.1571〕"),
("First, please proceed to Gate 3.","まず、3番ゲートへお進みください。〔No.1572〕"),
("navigate by the stars.","星によって進路を決める。〔No.1573〕"),
("roam around the world.","世界を歩き回る。〔No.1574〕"),
("flee to neighboring countries.","隣国へ逃亡する。〔No.1575〕"),
("When he saw me, he retreated into the room.","彼は私を見ると部屋に入って行った（部屋の中に退いた）。〔No.1576〕"),
("crawl into bed at four in the morning.","朝4時にベッドに這うように入る。〔No.1577〕"),
("a boat drifting on the waves.","波に漂う船。〔No.1578〕"),
("This does not alter the fact.","これによって事実が変わる（事実を変える）わけではない。〔No.1579〕"),
("convert the basement into an office.","地下室を事務所にする（転換する）。〔No.1580〕"),
("This shirt has shrunk.","このシャツは縮んだ。〔No.1581〕"),
("The flower withered in the heat.","暑さで花がしおれた。〔No.1582〕"),
("seasonal transitions.","季節の移り変わり。〔No.1583〕"),
("a distortion of the facts.","事実を歪めること。〔No.1584〕"),
("enhance my aesthetic sense.","美意識を磨く（向上させる）。〔No.1585〕"),
("reinforce the elbows of a jacket with patches.","当て布で上着のひじを補強する。〔No.1586〕"),
("foster good community relationships.","地域のよい関係を育てる。〔No.1587〕"),
("experience a family breakdown.","家庭の崩壊を経験する。〔No.1588〕"),
("collapse under the weight of the snow.","雪の重みで倒壊する。〔No.1589〕"),
("His remark rendered her speechless.","彼の発言に彼女は言葉を失った（彼女を話せない状態にした）。〔No.1590〕"),
("Alcohol impairs your ability to think.","アルコールは思考力を低下させる。〔No.1591〕"),
("Water quality is deteriorating rapidly.","水質が急速に悪化している。〔No.1592〕"),
("undermine Japan's social welfare system.","日本の社会福祉制度を揺るがす。〔No.1593〕"),
("tackle the unemployment problem.","失業問題に取り組む。〔No.1594〕"),
("have a lot of money at my disposal.","大金が自由に使える。〔No.1595〕"),
("We were confined to a small room.","私たちは小部屋に閉じ込められた。〔No.1596〕"),
("a designated smoking area.","指定された喫煙場所。〔No.1597〕"),
("regulate air pollution.","大気汚染を規制する。〔No.1598〕"),
("We are prohibited from holding second jobs.","私たちは副業を禁止されている。〔No.1599〕"),
("Please refrain from smoking.","たばこはお控えください。〔No.1600〕"),
("curb the spread of the disease.","病気の蔓延を抑制する。〔No.1601〕"),
("restrain my anger with difficulty.","やっとのことで怒りを抑える。〔No.1602〕"),
("halt the hunting of whales.","捕鯨を止める。〔No.1603〕"),
("animate the discussion.","議論を活気づける。〔No.1604〕"),
("spur economic growth.","経済成長を促す。〔No.1605〕"),
("urge him to rest.","彼に休むよう強く勧める。〔No.1606〕"),
("lure a deer into a trap.","シカを罠に誘い込む。〔No.1607〕"),
("Light stimulates plant growth.","光に刺激されて植物は育つ（光が植物の成長を刺激する）。〔No.1608〕"),
("feel compelled to resign because of the scandal.","スキャンダルのため辞職せざるを得ない（辞職を強いられている）と感じる。〔No.1609〕"),
("dictate a memo to my secretary.","秘書にメモを取らせる。〔No.1610〕"),
("Tears blurred my vision.","涙で視界がぼやけた（視界をぼやかした）。〔No.1611〕"),
("divert the course of a stream.","川の流れ（の進路）を変える。〔No.1612〕"),
("reverse the decision.","その決定を覆す。〔No.1613〕"),
("supplement my salary by working part-time.","アルバイトをして収入を補う。〔No.1614〕"),
("Pollution poses a threat to fish.","汚染が魚に脅威をもたらす。〔No.1615〕"),
("stroll along the riverbank.","土手を散策する。〔No.1616〕"),
("shrug my shoulders.","肩をすくめる。〔No.1617〕"),
("The dog is sniffing at the rug.","イヌが敷物の臭いをクンクンかいでいる。〔No.1618〕"),
("embrace his daughter.","彼の娘を抱擁する。〔No.1619〕"),
("betray her close friend.","彼女の親友を裏切る。〔No.1620〕"),
("I was deceived into believing that Kathy was a woman.","私はだまされて、キャシーは女性だと信じていた。〔No.1621〕"),
("Don't bully small children.","幼い子供たちをいじめるな。〔No.1622〕"),
("squeeze juice from an orange.","オレンジからジュースを搾る。〔No.1623〕"),
("insert a ticket into an automatic ticket gate.","自動改札機に切符を入れる。〔No.1624〕"),
("detach the hood from a jacket.","上着からフードを取り外す。〔No.1625〕"),
("withdraw a product from the market.","市場から製品を回収する。〔No.1626〕"),
("a building that can withstand major earthquakes.","大地震に耐えられる建物。〔No.1627〕"),
("exert strong leadership.","強いリーダーシップを発揮する。〔No.1628〕"),
("compile data for a report.","報告書のためにデータをまとめる。〔No.1629〕"),
("browse through the photographs.","写真にざっと目を通す。〔No.1630〕"),
("manipulate public opinion.","世論を操作する。〔No.1631〕"),
("implement political reform.","政治改革を実行する。〔No.1632〕"),
("execute the plan from Monday.","月曜日からその計画を実行する。〔No.1633〕"),
("host the Olympic Games.","オリンピックを主催する。〔No.1634〕"),
("a car mounted with a drive recorder.","ドライブレコーダーを搭載した車。〔No.1635〕"),
("discharge a black smoke.","黒い煙を出す。〔No.1636〕"),
("drain the bathtub.","風呂の水を抜く。〔No.1637〕"),
("get soaked in the rain.","雨でびしょ濡れになる。〔No.1638〕"),
("a lamp suspended from the ceiling.","天井からつるされているランプ。〔No.1639〕"),
("extract oil from olives.","オリーブの実から油を抽出する。〔No.1640〕"),
("The ambassador was immediately summoned home.","その大使は即座に本国に召喚された。〔No.1641〕"),
("embark on a huge project.","巨大なプロジェクトに着手する。〔No.1642〕"),
("thrust piles of banknotes into the bag.","かばんに札束を詰め込む。〔No.1643〕"),
("penetrate new markets.","新たな市場に入り込む。〔No.1644〕"),
("Don't intrude on her privacy.","彼女のプライバシーに立ち入るな。〔No.1645〕"),
("evade taxes.","税金を逃れる。〔No.1646〕"),
("Solar energy can be utilized for various purposes.","太陽エネルギーは様々な目的で利用できる。〔No.1647〕"),
("My head is spinning with a fever.","熱で頭がくらくらする。〔No.1648〕"),
("The helicopter plunged into the ocean.","そのヘリコプターは海に突っ込んだ。〔No.1649〕"),
("a rattling old bus.","ガタガタ音がする古いバス。〔No.1650〕"),
("The magician vanished without trace.","そのマジシャンは跡形もなく消えた。〔No.1651〕"),
("East Germany ceased to exist in 1990.","東ドイツは1990年に存在しなくなった。〔No.1652〕"),
("Haste makes waste.","慌てる（急ぐ）と損をする。〔No.1653〕"),
("The students concealed the facts from the teacher.","生徒たちは先生にその事実を隠した。〔No.1654〕"),
("disguise my sorrow with a smile.","笑顔で悲しみを隠す。〔No.1655〕"),
("That day coincided with my birthday.","その日は私の誕生日と重なった（同時に起きた）。〔No.1656〕"),
("Buddhism prevails throughout the country.","その国では仏教が隅々にまで普及している。〔No.1657〕"),
("sustain life on Earth.","地球上の生命を支える。〔No.1658〕"),
("A strong smell lingered in the elevator.","エレベーターの中に強烈な臭いが残っていた。〔No.1659〕"),
("The economy is beginning to revive after the slump.","経済は不況以降もち直し始めている。〔No.1660〕"),
("resume work.","仕事を再開する。〔No.1661〕"),
("inspect the elevators regularly.","定期的にエレベーターを検査する。〔No.1662〕"),
("investigate the ecology of eels.","ウナギの生態を調査する。〔No.1663〕"),
("detect a small amount of drug.","微量の麻薬を検出する。〔No.1664〕"),
("discern the slight difference.","そのささいな違いを見分ける。〔No.1665〕"),
("an expedition to the North Pole.","北極点への遠征。〔No.1666〕"),
("enroll in yoga classes.","ヨガのクラスに入る。〔No.1667〕"),
("discipline children.","子供をしつける。〔No.1668〕"),
("cram for the exam.","試験のために詰め込み勉強をする。〔No.1669〕"),
("the faculty of hearing.","聴力。〔No.1670〕"),
("pay my tuition by bank transfer.","授業料を銀行振込で払う。〔No.1671〕"),
("receive my diploma.","卒業証書を受け取る。〔No.1672〕"),
("mentor new recruits.","新入社員を指導する。〔No.1673〕"),
("medical ethics.","医療倫理。〔No.1674〕"),
("mysteries in astronomy.","天文学における謎。〔No.1675〕"),
("The archaeologist dug up an ancient pot.","その考古学者が古代のつぼを掘り出した。〔No.1676〕"),
("anthropologists studying about Neanderthals.","ネアンデルタール人を研究する人類学者。〔No.1677〕"),
("perceive the discovery as a major breakthrough.","その発見を飛躍的進歩と認識する。〔No.1678〕"),
("grasp the meaning of life.","人生の意味を理解する。〔No.1679〕"),
("conceive a good idea.","良いアイデアを思いつく。〔No.1680〕"),
("anticipate customers' needs.","顧客の要求を予想する。〔No.1681〕"),
("No one can foresee the future.","誰にも未来を予知することはできない。〔No.1682〕"),
("speculate about the future.","未来のことを推測する。〔No.1683〕"),
("infer the meaning from the context.","文脈から意味を推測する。〔No.1684〕"),
("What do you deduce from these figures?","これらの数字から何を推測しますか。〔No.1685〕"),
("ponder ways to bring more tourists to the town.","より多くの観光客を町に招致する方法を熟考する。〔No.1686〕"),
("contemplate marrying her.","彼女との結婚をじっくり考える。〔No.1687〕"),
("The doctor assured me that her life was not in danger.","医者は彼女の命に危険はないと私に保証してくれた。〔No.1688〕"),
("assess the method's efficiency.","その方法の効率を査定する。〔No.1689〕"),
("evaluate the employees' performances.","従業員の仕事ぶりを（正しく）評価する。〔No.1690〕"),
("verify the test results.","実験結果を検証する。〔No.1691〕"),
("compromise with him on this.","この点で彼と妥協する。〔No.1692〕"),
("reconcile an ideal with reality.","理想と現実を調和させる。〔No.1693〕"),
("We are alert to any possible danger.","私達はどんな危険にも対処できるように用心している。〔No.1694〕"),
("racial stereotypes.","人種による固定観念。〔No.1695〕"),
("gain a fresh perspective.","新たな視点を得る。〔No.1696〕"),
("I was both excited and worried at the prospect of becoming a father.","自分が父親になることを考える（見通す）と、わくわくすると同時に不安だった。〔No.1697〕"),
("make a vow never to smoke again.","二度とたばこは吸わないと誓う（誓いをする）。〔No.1698〕"),
("rely on my intuition.","自分の直感に頼る。〔No.1699〕"),
("Jessi is under the illusion that he loves her.","ジェシーは彼が自分のことを愛しているという幻想を抱いている。〔No.1700〕"),
("the criterion for judging wine.","ワインを判断する基準。〔No.1701〕"),
("tolerate unfair treatment.","不当な扱いに耐える。〔No.1702〕"),
("overlook some important evidence.","ある大切な証拠を見落とす。〔No.1703〕"),
("Japan's consumption of grain.","日本の穀物消費量。〔No.1704〕"),
("The village was hit by a severe drought.","その村は厳しい干ばつに見舞われた。〔No.1705〕"),
("water for irrigation.","灌漑用水。〔No.1706〕"),
("The sunset was quite a spectacle.","夕日が実に素晴らしい眺めだった。〔No.1707〕"),
("at an altitude of 10,000 meters.","高度10,000メートルで。〔No.1708〕"),
("natural catastrophes such as earthquakes, floods, and droughts.","地震、洪水、干ばつのような自然の大災害。〔No.1709〕"),
("develop tooth decay.","虫歯（歯の腐敗）になる。〔No.1710〕"),
("the erosion of the coastline.","その海岸線の浸食。〔No.1711〕"),
("a candle flame.","ろうそくの火。〔No.1712〕"),
("the Andromeda galaxy.","アンドロメダ銀河。〔No.1713〕"),
("A meteor shot across the night sky.","流星が夜空を横切った。〔No.1714〕"),
("I got a chill from standing outside.","外に立っていたので、冷えた。〔No.1715〕"),
("run through the meadow.","牧草地を走り抜ける。〔No.1716〕"),
("exploit natural resources.","天然資源を利用する。〔No.1717〕"),
("fish contaminated with chemicals.","化学物質で汚染された魚。〔No.1718〕"),
("erupt at regular intervals.","周期的に噴火する。〔No.1719〕"),
("Water evaporates when it is heated.","水は熱せられると蒸発する。〔No.1720〕"),
("fertile soil.","肥沃な土壌。〔No.1721〕"),
("The lion roared.","ライオンが吠えた。〔No.1722〕"),
("birds that inhabit New Zealand.","ニュージーランドに生息する鳥。〔No.1723〕"),
("nurture new industries.","新たな産業を育む。〔No.1724〕"),
("Owls sleep by day and hunt by night.","フクロウは昼に眠って、夜に狩りをする。〔No.1725〕"),
("Birds learn to fly by instinct.","鳥は本能的に飛び方を習得する。〔No.1726〕"),
("Zebras sometimes fall prey to lions.","シマウマはライオンの餌食になることがある。〔No.1727〕"),
("a timber dealer.","材木商。〔No.1728〕"),
("sweat-absorbing fabric.","汗を吸収しやすい布地。〔No.1729〕"),
("weave textiles.","織物を織る。〔No.1730〕"),
("weave fabric from wool.","羊毛から布地を織る。〔No.1731〕"),
("a competent translator.","有能な翻訳者。〔No.1732〕"),
("have an optimistic view of his future.","彼の将来を楽観する（楽観的な考えをもつ）。〔No.1733〕"),
("People feel more patriotic when they are abroad.","人は海外に出ると愛国的になるものだ。〔No.1734〕"),
("Sam is really naughty, but at the same time he is so cute.","サムは本当にいたずらだけど、同時にかわいいんだよね。〔No.1735〕"),
("a vigorous supporter.","精力的な支持者。〔No.1736〕"),
("long solitary walks.","一人でする（孤高の）何時間もの散歩。〔No.1737〕"),
("Emily is beautiful, but she is a little vulgar in her manner.","エミリーはとてもきれいだけれど、ふるまいは少し品がない。〔No.1738〕"),
("Scientists should be skeptical.","科学者は懐疑的であるべきだ。〔No.1739〕"),
("I am still haunted by regret.","まだ後悔している（後悔につきまとわれている）。〔No.1740〕"),
("Lisa is obsessed with her weight.","リサは体重のことばかり気にしている（体重のことで頭が一杯だ）。〔No.1741〕"),
("I was intent on my work.","仕事に没頭していた。〔No.1742〕"),
("indulge in drinking.","飲酒にふける。〔No.1743〕"),
("cling to his arm.","彼の腕にしがみつく。〔No.1744〕"),
("The company recalled cars that had a flaw in their braking system.","その会社はブレーキ系統に欠陥のある車をリコールした。〔No.1745〕"),
("a defect in that product.","その製品の欠陥。〔No.1746〕"),
("Skin color is a genetic trait.","肌の色は遺伝特性だ。〔No.1747〕"),
("draw an analogy between the brain and a computer.","脳をコンピュータに例える（脳とコンピュータの類似点を示す）。〔No.1748〕"),
("have no parallel in the 20th century.","20世紀には類似するものがない。〔No.1749〕"),
("play a crucial role in the negotiations.","その交渉で（極めて）重要な役割を演じる。〔No.1750〕"),
("My primary concern is your well-being.","私の今一番の関心は君の幸福だ。〔No.1751〕"),
("an integral part of our lives.","私たちの生活の不可欠な一部。〔No.1752〕"),
("Calcium is vital for healthy bones.","カルシウムは健康な骨に不可欠だ。〔No.1753〕"),
("have a profound impact on humans.","人間に深い影響を与える。〔No.1754〕"),
("face a grim reality.","厳しい現実に直面する。〔No.1755〕"),
("fierce competition.","激しい競争。〔No.1756〕"),
("Alaska's harsh climate.","アラスカの厳しい気候。〔No.1757〕"),
("a savage attack on the government.","政府への容赦ない攻撃。〔No.1758〕"),
("an acute problem.","深刻な問題。〔No.1759〕"),
("an ambiguous reply.","（どっちともとれる）曖昧な返事。〔No.1760〕"),
("an obscure problem.","（聞いたことがなく）よくわからない問題。〔No.1761〕"),
("a misleading article.","誤解を招く記事。〔No.1762〕"),
("Installing this program is quite straightforward.","このプログラムのインストールはかなりわかりやすい。〔No.1763〕"),
("explicit instructions.","明確な指示。〔No.1764〕"),
("a sophisticated alarm system.","高度な警報システム。〔No.1765〕"),
("an ingenious device to get rid of fleas.","ノミを駆除する独創的な装置。〔No.1766〕"),
("I'm paid a decent salary.","私はまずまずの給料をもらっている。〔No.1767〕"),
("The IT industry is thriving.","IT業界は好調だ。〔No.1768〕"),
("a man of great wit.","機知に富んだ男性。〔No.1769〕"),
("grow up in an affluent area.","裕福な地域で育つ。〔No.1770〕"),
("an elaborate meal.","手の込んだ料理。〔No.1771〕"),
("a prompt reply.","即答（迅速な返答）。〔No.1772〕"),
("a swift recovery.","素早い回復。〔No.1773〕"),
("The pen is mightier than the sword.","ペンは剣よりも強し。〔No.1774〕"),
("a restaurant renowned for its deep-fried chicken.","鶏のから揚げで有名な店。〔No.1775〕"),
("a dim light.","薄暗い明り。〔No.1776〕"),
("a gloomy old room.","薄暗く古びた部屋。〔No.1777〕"),
("a barren, rocky mountain.","草木も生えない（不毛）な岩山。〔No.1778〕"),
("sit idle in the morning.","午前中何もせずに座っている。〔No.1779〕"),
("a tiresome lecture.","退屈な講義。〔No.1780〕"),
("I got a shock from static electricity.","静電気がピリッときた。〔No.1781〕"),
("I am clumsy with my hands.","手先が不器用だ。〔No.1782〕"),
("Doris is ignorant of the world.","ドリスは世間知らずだ（世間について無知だ）。〔No.1783〕"),
("reckless driving.","無謀な運転。〔No.1784〕"),
("make an arbitrary decision.","独断的な決定をする。〔No.1785〕"),
("a foul smell.","不快な臭い。〔No.1786〕"),
("a shabby jacket.","みすぼらしい上着。〔No.1787〕"),
("Life in this town is monotonous.","この町での生活は単調だ。〔No.1788〕"),
("a crude rope bridge.","粗末なつり橋。〔No.1789〕"),
("fragile china.","壊れやすい陶磁器。〔No.1790〕"),
("Paper is vulnerable to moisture.","紙は湿気に弱い。〔No.1791〕"),
("The bicycle was covered in rust.","その自転車はさびついていた（さびで覆われていた）。〔No.1792〕"),
("the plight of children living in poverty.","貧困に生きる子供の窮状。〔No.1793〕"),
("conventional weapons.","通常兵器（従来の兵器）。〔No.1794〕"),
("an innate ability to make people laugh.","人を笑わせる先天的な才能。〔No.1795〕"),
("exotic plants.","外来植物。〔No.1796〕"),
("a tame dog.","おとなしい（飼いならされた）イヌ。〔No.1797〕"),
("a solemn ceremony.","厳粛な儀式。〔No.1798〕"),
("a manifest error.","明白な間違い。〔No.1799〕"),
("a dense fog.","濃い霧。〔No.1800〕"),
("a superficial knowledge of Japanese history.","日本史に関する表面的な知識。〔No.1801〕"),
("transparent glass.","透明なガラス。〔No.1802〕"),
("His words are not consistent with his behavior.","彼の発言は、行動と一致していない。〔No.1803〕"),
("give a coherent account of the event.","その出来事の一貫した説明をする。〔No.1804〕"),
('"Time is money" is a universal truth.',"「時は金なり」は普遍的な真理だ。〔No.1805〕"),
("walk upright.","直立歩行する（直立して歩く）。〔No.1806〕"),
("the underlying cause of the problem.","その問題の根底にある原因。〔No.1807〕"),
("a spontaneous cheer from the crowd.","群衆から自然に起こる声援。〔No.1808〕"),
("a mock exam.","模擬試験。〔No.1809〕"),
("Meg is liable to make mistakes.","メグはミスをしやすい。〔No.1810〕"),
("The converse is also true.","逆もまた真なり。〔No.1811〕"),
("preliminary market research.","事前の市場調査。〔No.1812〕"),
("his own version of the episode.","その出来事についての彼自身の解釈。〔No.1813〕"),
("a people indigenous to Australia.","オーストラリア固有の民族。〔No.1814〕"),
("Educational standards are declining.","教育水準が下がっている。〔No.1815〕"),
("multiply our profits.","私たちの利益を増やす。〔No.1816〕"),
("Land prices have soared in the last few years.","ここ数年で地価が急騰した。〔No.1817〕"),
("accumulate knowledge.","知識を蓄積する。〔No.1818〕"),
("exceed the speed limit.","制限速度を上回る。〔No.1819〕"),
("Canada is abundant in natural resources.","カナダは天然資源が豊富だ（天然資源において豊かだ）。〔No.1820〕"),
("give him ample opportunity to succeed.","成功する十分な機会を彼に与える。〔No.1821〕"),
("the gross national product (GNP).","国民総生産。〔No.1822〕"),
("an infinite number of stars.","無数の星。〔No.1823〕"),
("the sole survivor of the accident.","その事故の唯一の生存者。〔No.1824〕"),
("gather fallen leaves into a heap.","落ち葉を集めて山にする。〔No.1825〕"),
("a fraction of the cost.","その費用のほんの一部。〔No.1826〕"),
("work as a simultaneous interpreter.","同時通訳を務める。〔No.1827〕"),
("in medieval times.","中世（の時代）に。〔No.1828〕"),
("for decades.","数十年間。〔No.1829〕"),
("We left home at dawn.","私たちは夜明けに家を出た。〔No.1830〕"),
("over a span of five years.","５年という期間に。〔No.1831〕"),
("When does this license expire?","この免許はいつ期限が切れますか。〔No.1832〕"),
("postpone the meeting until tomorrow.","明日までの会議を延期する。〔No.1833〕"),
("These LED lights are equipped with artificial intelligence.","これらのLED証明には人工知能が搭載されている。〔No.1834〕"),
("transmit news all around the world.","世界中のニュースを伝える。〔No.1835〕"),
("He is endowed with a sense of humor.","彼にはユーモアのセンスがある（彼はユーモアのセンスを授けられている）。〔No.1836〕"),
("inherit the land from my grandfather.","祖父からその土地を受け継ぐ。〔No.1837〕"),
("retrieve the flight recorder.","フライトレコーダーを回収する。〔No.1838〕"),
("Natural gas has displaced coal.","天然ガスが石炭に取って代わった。〔No.1839〕"),
("the bond between mother and child.","母と子のきずな。〔No.1840〕"),
("Christmas tree ornaments.","クリスマスツリーの飾り。〔No.1841〕"),
("a birth certificate.","出生証明書。〔No.1842〕"),
("student accommodations.","学生用の宿泊施設。〔No.1843〕"),
("Turn left at the intersection.","交差点を左に曲がってください。〔No.1844〕"),
("bicycles left on the pavement.","歩道上に放置された自転車。〔No.1845〕"),
("a pedestrian bridge.","歩道橋（歩道者用の橋）。〔No.1846〕"),
("a school excursion to a zoo.","動物園への遠足。〔No.1847〕"),
("This itinerary looks too tight.","この旅程はタイトすぎるようだ。〔No.1848〕"),
("as a token of my thanks.","私の感謝のしるしとして。〔No.1849〕"),
("wash with laundry detergent.","洗濯洗剤で洗う。〔No.1850〕"),
("purchase an apartment in N.Y.","ニューヨークでマンションを購入する。〔No.1851〕"),
("dwell in a cave.","洞窟に住む。〔No.1852〕"),
("juvenile delinquency.","青少年の非行。〔No.1853〕"),
("a shy adolescent.","内気な思春期の若者。〔No.1854〕"),
("the initial stage of the disease.","その病気の初期（最初の）段階。〔No.1855〕"),
("In English, the subject precedes the verb.","英語では主語と動詞は前に来る。〔No.1856〕"),
("four phases in the lifecycle of a butterfly.","チョウの一生の４つの段階。〔No.1857〕"),
("Safety is our highest priority.","安全が最優先だ。〔No.1858〕"),
("the premise that everyone is equal.","皆平等であるという前提。〔No.1859〕"),
("according to legend.","伝説によると。〔No.1860〕"),
("a masterpiece by Leonardo da Vinci.","レオナルド・ダ・ヴィンチの傑作。〔No.1861〕"),
("perform a ritual.","儀式を行う。〔No.1862〕"),
("an object of worship.","崇拝の対象。〔No.1863〕"),
("the advent of new technology.","新技術の到来。〔No.1864〕"),
("swear on the Bible.","聖書に（手をのせて）誓う。〔No.1865〕"),
("confess to the crime.","犯罪を白状する。〔No.1866〕"),
("verbal communication.","言葉による意思の疎通。〔No.1867〕"),
("a 15-minute oral exam.","15分間の口述試験。〔No.1868〕"),
("make an eloquent speech.","雄弁な演説をする。〔No.1869〕"),
("a child's linguistic ability.","子どもの言語能力。〔No.1870〕"),
("the plot of the movie.","その映画の筋。〔No.1871〕"),
("write a draft of my speech.","スピーチの原稿を書く。〔No.1872〕"),
("These statements are true.","これらの記述は正しい。〔No.1873〕"),
('"Make haste slowly" is a kind of paradox.',"「急がば回れ」は一種の逆説だ。〔No.1874〕"),
('"Cool" is a slang for "nice."','"cool"は"nice"の俗語だ。〔No.1875〕'),
("Can you solve this riddle?","このなぞなぞがわかるかな。〔No.1876〕"),
("Let me clarify my point.","言いたい事を明確にします。〔No.1877〕"),
("posters exhibited in the hall.","ホールに展示されているポスター。〔No.1878〕"),
("portray the character as a selfish person.","利己的な人間としてその人物を描く。〔No.1879〕"),
("attribute his longevity to a glass of wine every day.","彼の長生きは毎日１杯のワインのおかげだと考える。〔No.1880〕"),
("console a grieving friend.","嘆き悲しんでいる友だちを慰める。〔No.1881〕"),
("The title aroused our interest.","そのタイトルが私たちの興味をかき立てた。〔No.1882〕"),
("I was distracted by someone calling my name.","誰かが私の名前を呼ぶ声に気をそらされた。〔No.1883〕"),
("plead with him to stay.","彼に留まるように懇願する。〔No.1884〕"),
("yearn for freedom.","自由を切望する。〔No.1885〕"),
("long for peace.","平和を切望する。〔No.1886〕"),
("Matt adores his grandchildren.","マットは孫を溺愛している。〔No.1887〕"),
("utter a cry of shock.","びっくりして叫び声を発する。〔No.1888〕"),
('"Oh my God!" she exclaimed.',"「なんてことなの！」と彼女は叫んだ。〔No.1889〕"),
("lament the death of the leader.","指導者の死を深く悲しむ。〔No.1890〕"),
("shed tears.","涙を流す。〔No.1891〕"),
("resent my boss's rude behavior.","上司の無礼なふるまいに憤慨する。〔No.1892〕"),
("dread going to the dentist.","歯医者に行くことを恐れる。〔No.1893〕"),
("I was startled by his voice.","私は彼の声で驚いた。〔No.1894〕"),
("Her performance thrilled me.","彼女の演奏は私をぞくぞくさせた。〔No.1895〕"),
("He humiliated me in front of my friends.","彼は友人の前で私に恥をかかせた。〔No.1896〕"),
("blush and look away.","顔を赤らめ、視線をそらす。〔No.1897〕"),
("frown on his behavior.","彼の行動に眉をひそめる。〔No.1898〕"),
("I was dismayed at the cost of the repair.","修理代に狼狽した（狼狽させられた）。〔No.1899〕"),
("I was bewildered by endless paperwork.","終わりのない事務処理に当惑した。〔No.1900〕"),
("Her strange habit perplexed me.","彼女の奇妙な癖は私を困惑させた。〔No.1901〕"),
("His attitudes always disgust me.","彼の態度はいつも私をむかつかせる。〔No.1902〕"),
("despise the neighbors because they are always gossiping.","うわさ話ばかりしているため、近所の人たちを軽蔑する。〔No.1903〕"),
("I was furious with him for leaving the baby alone in the car.","彼が赤ちゃんを車の中に放置したことに対して激怒した。〔No.1904〕"),
("an intimate conversation.","親密な会話。〔No.1905〕"),
("intriguing results.","興味深い結果。〔No.1906〕"),
("follow my conscience.","良心に従う。〔No.1907〕"),
("build my self-esteem.","自尊心を高める。〔No.1908〕"),
("beg for mercy.","慈悲を乞う。〔No.1909〕"),
("aspirations for peace.","平和への望み。〔No.1910〕"),
("Greg is in deep grief over her death.","グレッグは彼女の死に対して深い悲しみに暮れている。〔No.1911〕"),
("Henry is in great distress.","ヘンリーはひどく苦しい状況にいる。〔No.1912〕"),
("feel sick with apprehension.","不安で気分が悪い。〔No.1913〕"),
("Angie was lost in melancholy.","アンジーはふさぎ込んでいた。〔No.1914〕"),
("feel contempt for people obsessed with designer clothes.","ブランドの服に夢中になっている人々を軽蔑している。〔No.1915〕"),
("take her attitude as an insult.","彼女の態度を侮蔑と受け取る。〔No.1916〕"),
("a real nuisance.","本当に迷惑なもの。〔No.1917〕"),
("This drunk, violent man is a menace to those around him.","この酔って暴れる人は周りの者にとって厄介な存在だ。〔No.1918〕"),
("Anne was on the verge of tears.","アンは泣き出す寸前だった。〔No.1919〕"),
("an incentive to work hard.","一生懸命働く励み。〔No.1920〕"),
("predict the outcome of this election.","この選挙の結果を予測する。〔No.1921〕"),
("do a thorough cleaning of the house.","家の大掃除をする。〔No.1922〕"),
("His salary is adequate to support his family.","彼の給料は家族を養うには十分だ。〔No.1923〕"),
("an overall estimate.","全体的な見積もり。〔No.1924〕"),
("our ultimate goal.","私たちの最終的な目標。〔No.1925〕"),
("a genuine respect for Edison.","エジソンに対する心からの敬意。〔No.1926〕"),
("the need for radical reform.","抜本的な改革の必要性。〔No.1927〕"),
("take drastic measures.","抜本的な措置を講じる。〔No.1928〕"),
("Don't worry about such a trivial matter.","そんなささいなことで悩むな。〔No.1929〕"),
("That is virtually impossible.","それは事実上不可能だ。〔No.1930〕"),
("The car in front stopped abruptly.","前の車が不意に止まった。〔No.1931〕"),
("leave the letter on the desk deliberately.","手紙を故意に机の上に置く。〔No.1932〕"),
("The roads were icy; hence driving was not safe.","道路が凍っていた。それゆえ、車の運転は危険だった。〔No.1933〕"),
("two basic problems; namely, time and money.","２つの基本的な問題、すなわち時間とお金。〔No.1934〕"),
("Tom is quiet, whereas his brother is outgoing.","トムは無口だが一方、彼の弟はかなりおしゃべりだ。〔No.1935〕"),


       ]

import streamlit as st
from docx import Document
import io
# --- 🔐 パスワード認証機能 ---
# ここで「金庫の中のパスワード」と「入力されたパスワード」を照合します
password = st.text_input("パスワードを入力してください", type="password")
if password != st.secrets["MY_PASSWORD"]:
    st.warning("正しいパスワードを入力するとアプリが使えます。")
    st.stop()  # ここで処理を強制ストップ（これより下のコードは動きません）

# --- 2. 画面の設定 ---
st.title("単語・例文テスト作成アプリ 📝")
st.write("範囲（No.）を指定してボタンを押すと、テスト問題と解答を作成します。")

# 入力欄
col1, col2 = st.columns(2)
with col1:
    # 最初のデータのNo.に合わせて初期値を設定
    s = st.number_input('開始番号（No.）', min_value=1, value=1)
with col2:
    # データの個数を最大値に設定
    f = st.number_input('終了番号（No.）', min_value=1, value=len(leap1))

# --- 3. 作成ボタンが押されたら実行 ---
if st.button('テストを作成する！'):
    
    # エラーチェック
    if s > f:
        st.error("開始番号は終了番号より小さくしてください。")
        st.stop()
    if f > len(leap1):
        st.error(f"データは全部で {len(leap1)} 個しかありません。終了番号を減らしてください。")
        st.stop()

    # ドキュメントの準備
    leap_file = Document()       # 問題用
    leap_answer_file = Document() # 解答用
    
    # データの抽出（スライス）
    # リストは0番目から始まるので、開始位置は s-1 になります
    # Pythonのスライスは終了位置を含まないので、f はそのままでOKです
    target_data = leap1[s-1 : f] 
    
    # データがない場合
    if len(target_data) < 1:
        st.error("指定された範囲のデータが見つかりません。")
        st.stop()
        
    # シャッフル
    random.shuffle(target_data)
    
    # 辞書に変換（あなたのデータ形式ならこれで一発変換できます！）
    test_dict = dict(target_data)
    
    questions = list(test_dict.keys())   # 英語（左側）
    answers = list(test_dict.values())   # 日本語（右側）
    
    # ヘッダー作成
    header_text = "名前:＿＿＿＿＿＿＿＿＿＿＿＿＿＿, 範囲：No.{}～{}\n\n答えの〔No.～〕は単語番号です。\n".format(s, f)
    leap_file.add_paragraph(header_text)
    leap_answer_file.add_paragraph(header_text)

    # 問題作成ループ（指定した範囲の分だけ作る）
    # 元のコードにあった「20問」固定ではなく、選んだ範囲の個数分作ります
    for i in range(len(questions)):
        q_text = questions[i]
        a_text = answers[i]
        
        # 問題ファイル（英語＋下線）
        leap_file.add_paragraph(
            "Q{}:　{}\n{}".format(i + 1, q_text, '＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿')
        )
        # 解答ファイル（英語＋日本語答え）
        leap_answer_file.add_paragraph(
            "A{}:　{}\n{}".format(i + 1, q_text, a_text)
        )

    # --- 保存処理（メモリ上） ---
    bio_q = io.BytesIO()
    leap_file.save(bio_q)
    
    bio_a = io.BytesIO()
    leap_answer_file.save(bio_a)
    
    # --- ダウンロードボタン ---
    st.success(f"作成完了！範囲: {s}～{f} （全{len(questions)}問）")
    
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="📥 問題をダウンロード",
            data=bio_q.getvalue(),
            file_name=f"LEAP_テスト{s}～{f}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col2:
        st.download_button(
            label="📥 答えをダウンロード",
            data=bio_a.getvalue(),
            file_name=f"LEAP_答え{s}～{f}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )